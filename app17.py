import streamlit as st
import numpy as np
import faiss
import torch
import re
import html
from datetime import datetime
from io import BytesIO

from pypdf import PdfReader
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from sentence_transformers import SentenceTransformer
from transformers import AutoTokenizer, AutoModelForSeq2SeqLM
from deep_translator import GoogleTranslator


# -------------------------------
# PAGE CONFIG
# -------------------------------
st.set_page_config(
    page_title="Multilingual Document Assistant",
    layout="wide"
)


# -------------------------------
# PASTEL CLEAN UI CSS
# -------------------------------
st.markdown(
    """
    <style>
    .block-container {
        padding-top: 2rem;
        padding-bottom: 2rem;
        max-width: 1200px;
    }

    .hero-card {
        padding: 24px 26px;
        border-radius: 22px;
        background: linear-gradient(
            135deg,
            rgba(216, 180, 254, 0.22),
            rgba(191, 219, 254, 0.20),
            rgba(254, 215, 170, 0.16)
        );
        border: 1px solid rgba(148, 163, 184, 0.25);
        margin-bottom: 26px;
    }

    .main-title {
        font-size: 38px;
        font-weight: 800;
        color: inherit;
        margin-bottom: 6px;
        letter-spacing: -0.5px;
    }

    .subtitle {
        font-size: 16px;
        color: inherit;
        opacity: 0.72;
        margin-bottom: 8px;
    }

    .hero-small {
        font-size: 14px;
        opacity: 0.72;
        margin-top: 8px;
    }

    .mini-note {
        padding: 15px 18px;
        border-radius: 16px;
        background: linear-gradient(
            135deg,
            rgba(240, 249, 255, 0.16),
            rgba(236, 253, 245, 0.16)
        );
        border: 1px solid rgba(125, 211, 252, 0.25);
        color: inherit;
        margin-top: 10px;
        margin-bottom: 18px;
        line-height: 1.6;
    }

    .section-label {
        font-size: 20px;
        font-weight: 750;
        margin-bottom: 10px;
        color: inherit;
        letter-spacing: -0.2px;
    }

    .answer-box {
        padding: 22px;
        border-radius: 18px;
        background: linear-gradient(
            135deg,
            rgba(220, 252, 231, 0.18),
            rgba(219, 234, 254, 0.16)
        );
        border-left: 6px solid rgba(34, 197, 94, 0.85);
        border-top: 1px solid rgba(148, 163, 184, 0.20);
        border-right: 1px solid rgba(148, 163, 184, 0.20);
        border-bottom: 1px solid rgba(148, 163, 184, 0.20);
        color: inherit;
        margin-top: 12px;
        margin-bottom: 22px;
        line-height: 1.7;
        font-size: 16px;
        white-space: pre-wrap;
    }

    .summary-box {
        padding: 22px;
        border-radius: 18px;
        background: linear-gradient(
            135deg,
            rgba(254, 249, 195, 0.18),
            rgba(221, 214, 254, 0.16)
        );
        border-left: 6px solid rgba(234, 179, 8, 0.85);
        border-top: 1px solid rgba(148, 163, 184, 0.20);
        border-right: 1px solid rgba(148, 163, 184, 0.20);
        border-bottom: 1px solid rgba(148, 163, 184, 0.20);
        color: inherit;
        margin-top: 12px;
        margin-bottom: 22px;
        line-height: 1.7;
        font-size: 16px;
        white-space: pre-wrap;
    }

    div[data-testid="stMetric"] {
        background: linear-gradient(
            135deg,
            rgba(250, 245, 255, 0.18),
            rgba(239, 246, 255, 0.16)
        );
        border: 1px solid rgba(148, 163, 184, 0.24);
        padding: 17px;
        border-radius: 18px;
        box-shadow: 0 8px 22px rgba(15, 23, 42, 0.04);
    }

    div[data-testid="stMetric"] label {
        color: inherit !important;
        opacity: 0.78;
    }

    div[data-testid="stMetric"] div {
        color: inherit !important;
    }

    button[data-baseweb="tab"] {
        border-radius: 14px 14px 0 0;
        padding: 10px 16px;
        font-weight: 600;
    }

    button[data-baseweb="tab"][aria-selected="true"] {
        background-color: rgba(191, 219, 254, 0.18);
        border-bottom: 2px solid rgba(59, 130, 246, 0.85);
    }

    div.stButton > button {
        border-radius: 14px;
        border: 1px solid rgba(148, 163, 184, 0.28);
        background: linear-gradient(
            135deg,
            rgba(191, 219, 254, 0.24),
            rgba(221, 214, 254, 0.22)
        );
        color: inherit;
        font-weight: 650;
        padding: 0.55rem 1rem;
    }

    div.stButton > button:hover {
        border: 1px solid rgba(59, 130, 246, 0.55);
        transform: translateY(-1px);
        transition: 0.15s ease-in-out;
    }

    div.stDownloadButton > button {
        border-radius: 14px;
        background: linear-gradient(
            135deg,
            rgba(187, 247, 208, 0.24),
            rgba(191, 219, 254, 0.20)
        );
        border: 1px solid rgba(34, 197, 94, 0.30);
        color: inherit;
        font-weight: 650;
    }

    .stTextArea textarea {
        border-radius: 16px;
        line-height: 1.65;
        border: 1px solid rgba(148, 163, 184, 0.30);
        background-color: rgba(128, 128, 128, 0.06);
        color: inherit;
    }

    section[data-testid="stFileUploader"] {
        padding: 16px;
        border-radius: 18px;
        background-color: rgba(128, 128, 128, 0.05);
        border: 1px dashed rgba(148, 163, 184, 0.35);
    }

    div[data-testid="stExpander"] {
        border-radius: 16px;
        border: 1px solid rgba(148, 163, 184, 0.25);
        background-color: rgba(128, 128, 128, 0.04);
    }

    section[data-testid="stSidebar"] {
        background: linear-gradient(
            180deg,
            rgba(250, 245, 255, 0.18),
            rgba(239, 246, 255, 0.14)
        );
        border-right: 1px solid rgba(148, 163, 184, 0.18);
    }

    input, textarea {
        color: inherit !important;
    }
    </style>
    """,
    unsafe_allow_html=True
)


# -------------------------------
# LANGUAGE CODES
# -------------------------------
LANGUAGE_CODES = {
    "English": "en",
    "Tamil": "ta",
    "Hindi": "hi",
    "French": "fr",
    "Spanish": "es"
}


# -------------------------------
# REPORT LABELS
# -------------------------------
REPORT_LABELS = {
    "English": {
        "translation_report": "Document Translation Report",
        "qa_report": "Question Answering Report",
        "summary_report": "Document Summary Report",
        "generated_on": "Generated On",
        "source_file": "Source File",
        "translated_language": "Translated Language",
        "answer_language": "Answer Language",
        "summary_language": "Summary Language",
        "original_document": "Original Document",
        "translated_document": "Translated Document",
        "question": "Question",
        "answer": "Answer",
        "translated_evidence": "Translated Evidence",
        "original_evidence": "Original Evidence",
        "evidence": "Evidence",
        "source": "Source",
        "similarity_score": "Similarity Score",
        "summary": "Summary",
        "key_points": "Key Points"
    },
    "Tamil": {
        "translation_report": "ஆவண மொழிபெயர்ப்பு அறிக்கை",
        "qa_report": "கேள்வி பதில் அறிக்கை",
        "summary_report": "ஆவண சுருக்க அறிக்கை",
        "generated_on": "உருவாக்கப்பட்ட நேரம்",
        "source_file": "மூல கோப்பு",
        "translated_language": "மொழிபெயர்க்கப்பட்ட மொழி",
        "answer_language": "பதில் மொழி",
        "summary_language": "சுருக்க மொழி",
        "original_document": "அசல் ஆவணம்",
        "translated_document": "மொழிபெயர்க்கப்பட்ட ஆவணம்",
        "question": "கேள்வி",
        "answer": "பதில்",
        "translated_evidence": "மொழிபெயர்க்கப்பட்ட ஆதாரம்",
        "original_evidence": "அசல் ஆதாரம்",
        "evidence": "ஆதாரம்",
        "source": "மூலம்",
        "similarity_score": "ஒற்றுமை மதிப்பெண்",
        "summary": "சுருக்கம்",
        "key_points": "முக்கிய குறிப்புகள்"
    },
    "Hindi": {
        "translation_report": "दस्तावेज़ अनुवाद रिपोर्ट",
        "qa_report": "प्रश्न उत्तर रिपोर्ट",
        "summary_report": "दस्तावेज़ सारांश रिपोर्ट",
        "generated_on": "तैयार किया गया",
        "source_file": "स्रोत फ़ाइल",
        "translated_language": "अनुवादित भाषा",
        "answer_language": "उत्तर की भाषा",
        "summary_language": "सारांश की भाषा",
        "original_document": "मूल दस्तावेज़",
        "translated_document": "अनुवादित दस्तावेज़",
        "question": "प्रश्न",
        "answer": "उत्तर",
        "translated_evidence": "अनुवादित प्रमाण",
        "original_evidence": "मूल प्रमाण",
        "evidence": "प्रमाण",
        "source": "स्रोत",
        "similarity_score": "समानता स्कोर",
        "summary": "सारांश",
        "key_points": "मुख्य बिंदु"
    },
    "French": {
        "translation_report": "Rapport de traduction du document",
        "qa_report": "Rapport de questions-réponses",
        "summary_report": "Rapport de résumé du document",
        "generated_on": "Généré le",
        "source_file": "Fichier source",
        "translated_language": "Langue traduite",
        "answer_language": "Langue de réponse",
        "summary_language": "Langue du résumé",
        "original_document": "Document original",
        "translated_document": "Document traduit",
        "question": "Question",
        "answer": "Réponse",
        "translated_evidence": "Preuve traduite",
        "original_evidence": "Preuve originale",
        "evidence": "Preuve",
        "source": "Source",
        "similarity_score": "Score de similarité",
        "summary": "Résumé",
        "key_points": "Points clés"
    },
    "Spanish": {
        "translation_report": "Informe de traducción del documento",
        "qa_report": "Informe de preguntas y respuestas",
        "summary_report": "Informe de resumen del documento",
        "generated_on": "Generado el",
        "source_file": "Archivo fuente",
        "translated_language": "Idioma traducido",
        "answer_language": "Idioma de respuesta",
        "summary_language": "Idioma del resumen",
        "original_document": "Documento original",
        "translated_document": "Documento traducido",
        "question": "Pregunta",
        "answer": "Respuesta",
        "translated_evidence": "Evidencia traducida",
        "original_evidence": "Evidencia original",
        "evidence": "Evidencia",
        "source": "Fuente",
        "similarity_score": "Puntuación de similitud",
        "summary": "Resumen",
        "key_points": "Puntos clave"
    }
}


def label(key, language):
    return REPORT_LABELS.get(language, REPORT_LABELS["English"]).get(key, key)


# -------------------------------
# LOAD MODELS
# -------------------------------
@st.cache_resource
def load_embedding_model():
    return SentenceTransformer(
        "sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2"
    )


@st.cache_resource
def load_answer_model():
    model_name = "google/flan-t5-small"

    tokenizer = AutoTokenizer.from_pretrained(model_name)
    model = AutoModelForSeq2SeqLM.from_pretrained(model_name)

    device = "cuda" if torch.cuda.is_available() else "cpu"
    model = model.to(device)

    return tokenizer, model, device


embedding_model = load_embedding_model()
answer_tokenizer, answer_model, device = load_answer_model()


# -------------------------------
# TEXT CLEANING
# -------------------------------
def clean_text(text):
    text = text.replace("\n", " ")
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def split_into_sentences(text):
    sentences = re.split(r"(?<=[.!?।])\s+", text)
    return [sentence.strip() for sentence in sentences if sentence.strip()]


def clean_raw_text(text):
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    return text.strip()


def make_readable_paragraphs(text):
    text = clean_raw_text(text)

    text = re.sub(r"(\d+)\.([A-Za-z])", r"\1. \2", text)
    text = re.sub(r"\s+(\d+\.\s*)", r"\n\n\1", text)

    headings = [
        "Problem Statement",
        "Abstract",
        "Proposed Solution",
        "Tech Stack",
        "Workflow",
        "System Architecture",
        "Output",
        "Conclusion",
        "Features",
        "Advantages",
        "Limitations",
        "Future Enhancement",
        "Demo Scenario",
        "Modules",
        "Objective"
    ]

    for heading in headings:
        text = re.sub(
            rf"\s*({heading}\s*:?)",
            rf"\n\n\1\n",
            text,
            flags=re.IGNORECASE
        )

    blocks = re.split(r"\n\s*\n", text)
    final_blocks = []

    for block in blocks:
        block = re.sub(r"\s+", " ", block).strip()

        if not block:
            continue

        if len(block.split()) <= 8 and not block.endswith("."):
            final_blocks.append(block)
            continue

        sentences = split_into_sentences(block)

        if not sentences:
            final_blocks.append(block)
            continue

        current_para = ""

        for sentence in sentences:
            if len(current_para) + len(sentence) <= 700:
                if current_para:
                    current_para += " " + sentence
                else:
                    current_para = sentence
            else:
                if current_para:
                    final_blocks.append(current_para.strip())
                current_para = sentence

        if current_para:
            final_blocks.append(current_para.strip())

    return "\n\n".join(final_blocks).strip()


def fix_broken_pdf_lines(text):
    text = clean_raw_text(text)

    raw_lines = text.split("\n")
    lines = [line.strip() for line in raw_lines if line.strip()]

    if not lines:
        return ""

    small_line_count = 0

    for line in lines:
        if len(line.split()) <= 2:
            small_line_count += 1

    small_line_ratio = small_line_count / len(lines)

    if small_line_ratio > 0.45:
        merged_text = " ".join(lines)
        merged_text = re.sub(r"\s+", " ", merged_text).strip()
        return make_readable_paragraphs(merged_text)

    paragraphs = []
    current_para = ""

    for raw_line in raw_lines:
        line = raw_line.strip()

        if not line:
            if current_para.strip():
                paragraphs.append(current_para.strip())
                current_para = ""
            continue

        if not current_para:
            current_para = line
        else:
            if current_para.endswith("-"):
                current_para = current_para[:-1] + line
            else:
                current_para += " " + line

    if current_para.strip():
        paragraphs.append(current_para.strip())

    merged_text = "\n\n".join(paragraphs)

    return make_readable_paragraphs(merged_text)


def clean_for_docx(text):
    text = str(text)

    # Removes XML-invalid hidden control characters that can corrupt DOCX files.
    text = re.sub(r"[\x00-\x08\x0B\x0C\x0E-\x1F]", "", text)

    # Prevents extremely long single lines from becoming unstable in Word.
    text = text.replace("\u2028", "\n")
    text = text.replace("\u2029", "\n")

    return text


# -------------------------------
# TEXT EXTRACTION FROM FILE BYTES
# -------------------------------
def extract_pdf_text_for_qa(file_bytes):
    text = ""
    pdf_reader = PdfReader(BytesIO(file_bytes))

    for page in pdf_reader.pages:
        page_text = page.extract_text()

        if page_text:
            text += page_text + " "

    return clean_text(text)


def extract_docx_text_for_qa(file_bytes):
    text = ""
    document = Document(BytesIO(file_bytes))

    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            text += paragraph.text + " "

    return clean_text(text)


def extract_txt_text_for_qa(file_bytes):
    text = file_bytes.decode("utf-8", errors="ignore")
    return clean_text(text)


def extract_text_for_qa(file_name, file_bytes):
    file_name = file_name.lower()

    if file_name.endswith(".pdf"):
        return extract_pdf_text_for_qa(file_bytes)

    elif file_name.endswith(".docx"):
        return extract_docx_text_for_qa(file_bytes)

    elif file_name.endswith(".txt"):
        return extract_txt_text_for_qa(file_bytes)

    return ""


def extract_pdf_text_for_translation(file_bytes):
    page_texts = []
    pdf_reader = PdfReader(BytesIO(file_bytes))

    for page in pdf_reader.pages:
        page_text = page.extract_text()

        if page_text:
            page_texts.append(page_text)

    raw_text = "\n\n".join(page_texts)
    return fix_broken_pdf_lines(raw_text)


def extract_docx_text_for_translation(file_bytes):
    document = Document(BytesIO(file_bytes))
    paragraphs = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()

        if text:
            paragraphs.append(text)

    raw_text = "\n\n".join(paragraphs)
    return make_readable_paragraphs(raw_text)


def extract_txt_text_for_translation(file_bytes):
    raw_text = file_bytes.decode("utf-8", errors="ignore")
    return make_readable_paragraphs(raw_text)


def extract_text_for_translation(file_name, file_bytes):
    file_name = file_name.lower()

    if file_name.endswith(".pdf"):
        return extract_pdf_text_for_translation(file_bytes)

    elif file_name.endswith(".docx"):
        return extract_docx_text_for_translation(file_bytes)

    elif file_name.endswith(".txt"):
        return extract_txt_text_for_translation(file_bytes)

    return ""


# -------------------------------
# CHUNKING
# -------------------------------
def chunk_text(text, chunk_size=160, overlap=35):
    words = text.split()
    chunks = []

    start = 0

    while start < len(words):
        end = start + chunk_size
        chunk = " ".join(words[start:end])

        if chunk.strip():
            chunks.append(chunk.strip())

        start = end - overlap

    return chunks


def split_long_text_for_translation(text, max_chars=2800):
    if len(text) <= max_chars:
        return [text]

    sentences = split_into_sentences(text)

    if not sentences:
        return [text[i:i + max_chars] for i in range(0, len(text), max_chars)]

    parts = []
    current_part = ""

    for sentence in sentences:
        if len(sentence) > max_chars:
            if current_part.strip():
                parts.append(current_part.strip())
                current_part = ""

            for i in range(0, len(sentence), max_chars):
                parts.append(sentence[i:i + max_chars])

        elif len(current_part) + len(sentence) + 1 <= max_chars:
            if current_part:
                current_part += " " + sentence
            else:
                current_part = sentence

        else:
            if current_part.strip():
                parts.append(current_part.strip())

            current_part = sentence

    if current_part.strip():
        parts.append(current_part.strip())

    return parts


# -------------------------------
# TRANSLATION
# -------------------------------
def translate_to_english(text):
    try:
        return GoogleTranslator(
            source="auto",
            target="en"
        ).translate(text)

    except Exception:
        return text


def translate_text(text, target_language):
    target_code = LANGUAGE_CODES.get(target_language, "en")

    if target_code == "en":
        return text

    try:
        return GoogleTranslator(
            source="auto",
            target=target_code
        ).translate(text)

    except Exception:
        return text


def translate_full_document(text, target_language):
    target_code = LANGUAGE_CODES.get(target_language, "en")

    if target_code == "en":
        return text

    paragraphs = re.split(r"\n\s*\n", text)
    paragraphs = [para.strip() for para in paragraphs if para.strip()]

    translation_units = []

    for paragraph in paragraphs:
        parts = split_long_text_for_translation(paragraph)
        translation_units.append(parts)

    total_parts = sum(len(parts) for parts in translation_units)

    if total_parts == 0:
        return ""

    progress_bar = st.progress(0)
    status_text = st.empty()

    completed_parts = 0
    translated_paragraphs = []

    for parts in translation_units:
        translated_parts = []

        for part in parts:
            status_text.write(
                f"Translating section {completed_parts + 1} of {total_parts}..."
            )

            try:
                translated_part = GoogleTranslator(
                    source="auto",
                    target=target_code
                ).translate(part)

            except Exception:
                translated_part = part

            translated_parts.append(translated_part)

            completed_parts += 1
            progress_bar.progress(completed_parts / total_parts)

        translated_paragraph = " ".join(translated_parts)
        translated_paragraphs.append(translated_paragraph.strip())

    status_text.empty()

    return "\n\n".join(translated_paragraphs).strip()


# -------------------------------
# FAISS
# -------------------------------
def create_faiss_index(chunks):
    embeddings = embedding_model.encode(chunks, convert_to_numpy=True)
    embeddings = np.array(embeddings).astype("float32")

    faiss.normalize_L2(embeddings)

    dimension = embeddings.shape[1]
    index = faiss.IndexFlatIP(dimension)

    index.add(embeddings)

    return index


def search_query(query, index, chunks, chunk_sources, top_k=3):
    query_embedding = embedding_model.encode([query], convert_to_numpy=True)
    query_embedding = np.array(query_embedding).astype("float32")

    faiss.normalize_L2(query_embedding)

    scores, indices = index.search(query_embedding, top_k)

    results = []

    for score, idx in zip(scores[0], indices[0]):
        if idx != -1:
            results.append(
                {
                    "chunk": chunks[idx],
                    "source": chunk_sources[idx],
                    "score": float(score)
                }
            )

    return results


# -------------------------------
# SLM HELPER
# -------------------------------
def generate_with_slm(prompt, max_new_tokens=180):
    inputs = answer_tokenizer(
        prompt,
        return_tensors="pt",
        truncation=True,
        max_length=1024
    ).to(device)

    with torch.no_grad():
        outputs = answer_model.generate(
            **inputs,
            max_new_tokens=max_new_tokens,
            num_beams=4,
            no_repeat_ngram_size=3,
            early_stopping=True
        )

    result = answer_tokenizer.decode(
        outputs[0],
        skip_special_tokens=True
    )

    return result.strip()


# -------------------------------
# ANSWER GENERATION
# -------------------------------
def extractive_answer(query, chunks, max_sentences=4):
    all_sentences = []

    for chunk in chunks:
        sentences = split_into_sentences(chunk)
        all_sentences.extend(sentences)

    if not all_sentences:
        return "The answer is not available in the uploaded documents."

    sentence_embeddings = embedding_model.encode(
        all_sentences,
        convert_to_numpy=True
    )

    query_embedding = embedding_model.encode(
        [query],
        convert_to_numpy=True
    )

    sentence_embeddings = np.array(sentence_embeddings).astype("float32")
    query_embedding = np.array(query_embedding).astype("float32")

    faiss.normalize_L2(sentence_embeddings)
    faiss.normalize_L2(query_embedding)

    scores = np.dot(sentence_embeddings, query_embedding.T).flatten()
    top_indices = scores.argsort()[-max_sentences:][::-1]

    selected_sentences = []

    for idx in top_indices:
        sentence = all_sentences[idx]

        if sentence not in selected_sentences:
            selected_sentences.append(sentence)

    return " ".join(selected_sentences)


def generate_clean_answer(query, retrieved_chunks):
    english_query = translate_to_english(query)

    english_context_parts = []

    for chunk in retrieved_chunks[:3]:
        short_chunk = chunk[:3500]
        english_chunk = translate_to_english(short_chunk)
        english_context_parts.append(english_chunk)

    context = "\n\n".join(english_context_parts)

    prompt = f"""
Answer the question using only the given context.

Rules:
- Give a clear and simple answer.
- Do not copy the full context.
- Do not add outside information.
- If the answer is not found, say: The answer is not available in the uploaded documents.

Context:
{context}

Question:
{english_query}

Answer:
"""

    answer = generate_with_slm(prompt, max_new_tokens=180)

    weak_answers = [
        "",
        "answer:",
        "context:",
        "The answer is not available in the uploaded documents."
    ]

    if answer.lower() in [item.lower() for item in weak_answers] or len(answer.split()) < 8:
        answer = extractive_answer(
            query=english_query,
            chunks=english_context_parts
        )

    return answer


# -------------------------------
# SUMMARY GENERATION
# -------------------------------
def prepare_summary_text(text, max_chars=4500):
    paragraphs = re.split(r"\n\s*\n", text)
    paragraphs = [para.strip() for para in paragraphs if para.strip()]

    selected_text = ""

    for paragraph in paragraphs:
        if len(selected_text) + len(paragraph) + 2 <= max_chars:
            selected_text += paragraph + "\n\n"
        else:
            break

    if not selected_text.strip():
        selected_text = text[:max_chars]

    return selected_text.strip()


def fallback_summary(text):
    sentences = split_into_sentences(text)

    if not sentences:
        return "1. Summary is not available for this document."

    selected = sentences[:4]

    points = []

    for i, sentence in enumerate(selected, start=1):
        points.append(f"{i}. {sentence}")

    return "\n".join(points)


def fallback_key_points(text):
    sentences = split_into_sentences(text)
    selected = sentences[:5]

    if not selected:
        return "- Key points are not available for this document."

    points = []

    for sentence in selected:
        points.append(f"- {sentence}")

    return "\n".join(points)


def generate_document_summary(text, target_language):
    source_text = prepare_summary_text(text, max_chars=4500)
    english_text = translate_to_english(source_text)

    prompt = f"""
Read the following document and create a clear summary.

Rules:
- Write the summary as 4 numbered points.
- Each point should be one full sentence.
- Keep it simple and useful.
- Do not add outside information.
- Do not give only one line.

Document:
{english_text}

4-point summary:
"""

    summary = generate_with_slm(prompt, max_new_tokens=260)

    if not summary or len(summary.split()) < 35:
        summary = fallback_summary(english_text)

    if "1." not in summary and "- " not in summary:
        sentences = split_into_sentences(summary)

        if len(sentences) >= 3:
            numbered_points = []

            for i, sentence in enumerate(sentences[:4], start=1):
                numbered_points.append(f"{i}. {sentence}")

            summary = "\n".join(numbered_points)

        else:
            summary = fallback_summary(english_text)

    translated_summary = translate_text(
        text=summary,
        target_language=target_language
    )

    return translated_summary


def generate_document_key_points(text, target_language):
    source_text = prepare_summary_text(text, max_chars=4500)
    english_text = translate_to_english(source_text)

    prompt = f"""
Extract 5 important key points from the following document.
Write each point as a short bullet.
Do not add outside information.

Document:
{english_text}

Key points:
"""

    key_points = generate_with_slm(prompt, max_new_tokens=220)

    if not key_points or len(key_points.split()) < 10:
        key_points = fallback_key_points(english_text)

    translated_key_points = translate_text(
        text=key_points,
        target_language=target_language
    )

    return translated_key_points


# -------------------------------
# DOCX HELPERS
# -------------------------------
def set_run_font(run, size=11, bold=False):
    run.font.name = "Nirmala UI"
    run.font.size = Pt(size)
    run.font.bold = bold

    run_element = run._element
    run_properties = run_element.get_or_add_rPr()
    run_fonts = run_properties.get_or_add_rFonts()

    run_fonts.set(qn("w:ascii"), "Nirmala UI")
    run_fonts.set(qn("w:hAnsi"), "Nirmala UI")
    run_fonts.set(qn("w:eastAsia"), "Nirmala UI")
    run_fonts.set(qn("w:cs"), "Nirmala UI")


def split_long_line(line, limit=1800):
    if len(line) <= limit:
        return [line]

    parts = []

    for i in range(0, len(line), limit):
        parts.append(line[i:i + limit])

    return parts


def add_multiline_text(document, text):
    text = clean_for_docx(text)
    lines = text.split("\n")

    for line in lines:
        if not line.strip():
            document.add_paragraph("")
            continue

        for part in split_long_line(line.strip()):
            paragraph = document.add_paragraph()
            run = paragraph.add_run(part)
            set_run_font(run, size=11)


def create_docx_file(title, sections):
    document = Document()

    styles = document.styles
    styles["Normal"].font.name = "Nirmala UI"
    styles["Normal"].font.size = Pt(11)

    safe_title = clean_for_docx(title)

    heading = document.add_heading("", level=0)
    heading_run = heading.add_run(safe_title)
    set_run_font(heading_run, size=18, bold=True)

    generated_para = document.add_paragraph()
    generated_run = generated_para.add_run(
        f"Generated On: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
    )
    set_run_font(generated_run, size=10)

    for section_title, section_text in sections:
        safe_section_title = clean_for_docx(section_title)

        section_heading = document.add_heading("", level=1)
        section_run = section_heading.add_run(safe_section_title)
        set_run_font(section_run, size=14, bold=True)

        add_multiline_text(document, section_text)

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)

    return buffer.getvalue()


# -------------------------------
# REPORTS
# -------------------------------
def create_qa_report(query, preferred_language, final_answer, results):
    report = ""
    report += label("qa_report", preferred_language).upper() + "\n"
    report += "=" * 60 + "\n\n"

    report += f"{label('generated_on', preferred_language)}: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
    report += f"{label('answer_language', preferred_language)}: {preferred_language}\n\n"

    report += label("question", preferred_language).upper() + "\n"
    report += "-" * 60 + "\n"
    report += query + "\n\n"

    report += label("answer", preferred_language).upper() + "\n"
    report += "-" * 60 + "\n"
    report += final_answer + "\n\n"

    report += label("evidence", preferred_language).upper() + "\n"
    report += "-" * 60 + "\n"

    for i, result in enumerate(results, start=1):
        translated_evidence = translate_text(
            text=result["chunk"],
            target_language=preferred_language
        )

        report += f"\n{label('evidence', preferred_language)} {i}\n"
        report += f"{label('source', preferred_language)}: {result['source']}\n"
        report += f"{label('similarity_score', preferred_language)}: {result['score']:.4f}\n\n"

        report += label("translated_evidence", preferred_language) + "\n"
        report += translated_evidence + "\n\n"

        report += label("original_evidence", preferred_language) + "\n"
        report += result["chunk"] + "\n"

    return report


def create_qa_docx(query, preferred_language, final_answer, results):
    evidence_text = ""

    for i, result in enumerate(results, start=1):
        translated_evidence = translate_text(
            text=result["chunk"],
            target_language=preferred_language
        )

        evidence_text += f"{label('evidence', preferred_language)} {i}\n"
        evidence_text += f"{label('source', preferred_language)}: {result['source']}\n"
        evidence_text += f"{label('similarity_score', preferred_language)}: {result['score']:.4f}\n\n"
        evidence_text += f"{label('translated_evidence', preferred_language)}:\n"
        evidence_text += translated_evidence + "\n\n"
        evidence_text += f"{label('original_evidence', preferred_language)}:\n"
        evidence_text += result["chunk"] + "\n\n"

    sections = [
        (label("question", preferred_language), query),
        (label("answer", preferred_language), final_answer),
        (label("evidence", preferred_language), evidence_text)
    ]

    return create_docx_file(
        title=label("qa_report", preferred_language),
        sections=sections
    )


def create_translation_report(file_name, target_language, original_text, translated_text):
    report = ""
    report += label("translation_report", target_language).upper() + "\n"
    report += "=" * 60 + "\n\n"

    report += f"{label('generated_on', target_language)}: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
    report += f"{label('source_file', target_language)}: {file_name}\n"
    report += f"{label('translated_language', target_language)}: {target_language}\n\n"

    report += label("original_document", target_language).upper() + "\n"
    report += "-" * 60 + "\n\n"
    report += original_text + "\n\n"

    report += label("translated_document", target_language).upper() + "\n"
    report += "-" * 60 + "\n\n"
    report += translated_text

    return report


def create_translation_docx(file_name, target_language, original_text, translated_text):
    sections = [
        (label("source_file", target_language), file_name),
        (label("original_document", target_language), original_text),
        (label("translated_document", target_language), translated_text)
    ]

    return create_docx_file(
        title=label("translation_report", target_language),
        sections=sections
    )


def create_summary_report(file_name, target_language, summary, key_points):
    report = ""
    report += label("summary_report", target_language).upper() + "\n"
    report += "=" * 60 + "\n\n"

    report += f"{label('generated_on', target_language)}: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
    report += f"{label('source_file', target_language)}: {file_name}\n"
    report += f"{label('summary_language', target_language)}: {target_language}\n\n"

    report += label("summary", target_language).upper() + "\n"
    report += "-" * 60 + "\n\n"
    report += summary + "\n\n"

    report += label("key_points", target_language).upper() + "\n"
    report += "-" * 60 + "\n\n"
    report += key_points

    return report


def create_summary_docx(file_name, target_language, summary, key_points):
    sections = [
        (label("source_file", target_language), file_name),
        (label("summary", target_language), summary),
        (label("key_points", target_language), key_points)
    ]

    return create_docx_file(
        title=label("summary_report", target_language),
        sections=sections
    )


# -------------------------------
# SESSION STATE
# -------------------------------
if "uploaded_docs" not in st.session_state:
    st.session_state.uploaded_docs = {}

if "formatted_texts" not in st.session_state:
    st.session_state.formatted_texts = {}

if "chunks" not in st.session_state:
    st.session_state.chunks = []

if "chunk_sources" not in st.session_state:
    st.session_state.chunk_sources = []

if "index" not in st.session_state:
    st.session_state.index = None

if "file_names" not in st.session_state:
    st.session_state.file_names = []

if "query_history" not in st.session_state:
    st.session_state.query_history = []

if "latest_qa_report" not in st.session_state:
    st.session_state.latest_qa_report = None

if "latest_qa_docx" not in st.session_state:
    st.session_state.latest_qa_docx = None

if "translated_document" not in st.session_state:
    st.session_state.translated_document = None

if "translation_report" not in st.session_state:
    st.session_state.translation_report = None

if "translation_docx" not in st.session_state:
    st.session_state.translation_docx = None

if "selected_translation_file" not in st.session_state:
    st.session_state.selected_translation_file = None

if "translation_language" not in st.session_state:
    st.session_state.translation_language = None

if "summary_report" not in st.session_state:
    st.session_state.summary_report = None

if "summary_docx" not in st.session_state:
    st.session_state.summary_docx = None

if "document_summary" not in st.session_state:
    st.session_state.document_summary = None

if "document_key_points" not in st.session_state:
    st.session_state.document_key_points = None

if "summary_file" not in st.session_state:
    st.session_state.summary_file = None

if "summary_language" not in st.session_state:
    st.session_state.summary_language = None


# -------------------------------
# HEADER
# -------------------------------
st.markdown(
    """
    <div class="hero-card">
        <div class="main-title">Multilingual Document Assistant</div>
        <div class="subtitle">
            Upload your document once, translate it side-by-side, summarize it, or ask questions from it in your preferred language.
        </div>
        <div class="hero-small">
            A simple workspace for multilingual document reading, translation, summary, and evidence-based answers.
        </div>
    </div>
    """,
    unsafe_allow_html=True
)


# -------------------------------
# SIDEBAR
# -------------------------------
st.sidebar.title("Settings")

preferred_language = st.sidebar.selectbox(
    "Preferred Language",
    ["English", "Tamil", "Hindi", "French", "Spanish"]
)

st.sidebar.markdown("---")

if st.sidebar.button("Reset"):
    st.session_state.uploaded_docs = {}
    st.session_state.formatted_texts = {}
    st.session_state.chunks = []
    st.session_state.chunk_sources = []
    st.session_state.index = None
    st.session_state.file_names = []
    st.session_state.query_history = []
    st.session_state.latest_qa_report = None
    st.session_state.latest_qa_docx = None
    st.session_state.translated_document = None
    st.session_state.translation_report = None
    st.session_state.translation_docx = None
    st.session_state.selected_translation_file = None
    st.session_state.translation_language = None
    st.session_state.summary_report = None
    st.session_state.summary_docx = None
    st.session_state.document_summary = None
    st.session_state.document_key_points = None
    st.session_state.summary_file = None
    st.session_state.summary_language = None
    st.success("Reset completed.")


# -------------------------------
# DASHBOARD
# -------------------------------
col1, col2, col3 = st.columns(3)

with col1:
    st.metric("Documents", len(st.session_state.file_names))

with col2:
    st.metric("Text Sections", len(st.session_state.chunks))

with col3:
    status = "Ready" if st.session_state.index is not None else "Waiting"
    st.metric("Status", status)


# -------------------------------
# TABS
# -------------------------------
tab1, tab2, tab3, tab4, tab5 = st.tabs(
    [
        "Upload",
        "Translate Document",
        "Ask Questions",
        "Summarize",
        "History"
    ]
)


# -------------------------------
# TAB 1: UPLOAD
# -------------------------------
with tab1:
    st.markdown('<div class="section-label">Upload your documents</div>', unsafe_allow_html=True)

    st.markdown(
        """
        <div class="mini-note">
        Upload once. The same document will be used for translation, summary, and question answering.
        </div>
        """,
        unsafe_allow_html=True
    )

    uploaded_files = st.file_uploader(
        "Choose PDF, TXT, or DOCX files",
        type=["pdf", "txt", "docx"],
        accept_multiple_files=True,
        key="single_upload"
    )

    if uploaded_files:
        st.write("Selected files:")

        for uploaded_file in uploaded_files:
            st.write(f"- {uploaded_file.name}")

        if st.button("Process Documents"):
            uploaded_docs = {}
            formatted_texts = {}
            all_chunks = []
            chunk_sources = []
            file_names = []

            with st.spinner("Preparing documents..."):
                for uploaded_file in uploaded_files:
                    file_name = uploaded_file.name
                    file_bytes = uploaded_file.getvalue()

                    uploaded_docs[file_name] = file_bytes

                    qa_text = extract_text_for_qa(file_name, file_bytes)
                    formatted_text = extract_text_for_translation(file_name, file_bytes)

                    if formatted_text.strip():
                        formatted_texts[file_name] = formatted_text

                    if qa_text.strip():
                        chunks = chunk_text(qa_text)

                        for chunk in chunks:
                            all_chunks.append(chunk)
                            chunk_sources.append(file_name)

                        file_names.append(file_name)

                if all_chunks:
                    index = create_faiss_index(all_chunks)

                    st.session_state.uploaded_docs = uploaded_docs
                    st.session_state.formatted_texts = formatted_texts
                    st.session_state.chunks = all_chunks
                    st.session_state.chunk_sources = chunk_sources
                    st.session_state.index = index
                    st.session_state.file_names = file_names

                    st.session_state.translated_document = None
                    st.session_state.translation_report = None
                    st.session_state.translation_docx = None
                    st.session_state.selected_translation_file = None
                    st.session_state.translation_language = None

                    st.session_state.summary_report = None
                    st.session_state.summary_docx = None
                    st.session_state.document_summary = None
                    st.session_state.document_key_points = None
                    st.session_state.summary_file = None
                    st.session_state.summary_language = None

                    st.session_state.latest_qa_report = None
                    st.session_state.latest_qa_docx = None

                    st.success("Documents are ready.")

                else:
                    st.error("No readable text found in the uploaded documents.")

    if st.session_state.file_names:
        st.markdown("### Ready documents")
        for name in st.session_state.file_names:
            st.write(f"- {name}")


# -------------------------------
# TAB 2: TRANSLATE DOCUMENT
# -------------------------------
with tab2:
    st.markdown('<div class="section-label">Side-by-side document translation</div>', unsafe_allow_html=True)

    if not st.session_state.formatted_texts:
        st.warning("Please upload and process a document first.")

    else:
        selected_file = st.selectbox(
            "Select document",
            list(st.session_state.formatted_texts.keys()),
            key="translation_selectbox"
        )

        st.write(f"Translation language: **{preferred_language}**")

        if st.button("Translate Document"):
            original_text = st.session_state.formatted_texts[selected_file]

            with st.spinner("Translating document..."):
                translated_text = translate_full_document(
                    text=original_text,
                    target_language=preferred_language
                )

            st.session_state.translated_document = translated_text
            st.session_state.selected_translation_file = selected_file
            st.session_state.translation_language = preferred_language
            st.session_state.translation_report = create_translation_report(
                file_name=selected_file,
                target_language=preferred_language,
                original_text=original_text,
                translated_text=translated_text
            )
            st.session_state.translation_docx = create_translation_docx(
                file_name=selected_file,
                target_language=preferred_language,
                original_text=original_text,
                translated_text=translated_text
            )

            st.success("Translation completed.")

        if st.session_state.translated_document:
            original_text = st.session_state.formatted_texts[
                st.session_state.selected_translation_file
            ]

            st.markdown("### Document View")

            left_col, right_col = st.columns(2)

            with left_col:
                st.markdown("**Original Document**")
                st.text_area(
                    "Original",
                    value=original_text[:7000],
                    height=520,
                    label_visibility="collapsed"
                )

            with right_col:
                lang = st.session_state.translation_language or preferred_language
                st.markdown(f"**Translated Document ({lang})**")
                st.text_area(
                    "Translated",
                    value=st.session_state.translated_document[:7000],
                    height=520,
                    label_visibility="collapsed"
                )

            if len(st.session_state.translated_document) > 7000:
                st.info("Preview shows the first 7000 characters. Download the complete translation below.")

            download_col1, download_col2 = st.columns(2)

            with download_col1:
                st.download_button(
                    label="Download Translation TXT",
                    data=st.session_state.translation_report,
                    file_name="translated_document_report.txt",
                    mime="text/plain"
                )

            with download_col2:
                st.download_button(
                    label="Download Translation DOCX",
                    data=st.session_state.translation_docx,
                    file_name="translated_document_report.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )


# -------------------------------
# TAB 3: ASK QUESTIONS
# -------------------------------
with tab3:
    st.markdown('<div class="section-label">Ask questions from your document</div>', unsafe_allow_html=True)

    if not st.session_state.chunks or st.session_state.index is None:
        st.warning("Please upload and process a document first.")

    else:
        query = st.text_input(
            "Ask your question",
            placeholder="Example: What is this document about?"
        )

        if st.button("Get Answer"):
            if not query.strip():
                st.error("Please enter a question.")

            else:
                with st.spinner("Searching document..."):
                    results = search_query(
                        query=query,
                        index=st.session_state.index,
                        chunks=st.session_state.chunks,
                        chunk_sources=st.session_state.chunk_sources,
                        top_k=3
                    )

                retrieved_chunks = [result["chunk"] for result in results]

                with st.spinner("Writing answer..."):
                    english_answer = generate_clean_answer(
                        query=query,
                        retrieved_chunks=retrieved_chunks
                    )

                    final_answer = translate_text(
                        text=english_answer,
                        target_language=preferred_language
                    )

                st.session_state.query_history.append(
                    {
                        "time": datetime.now().strftime("%H:%M:%S"),
                        "question": query,
                        "language": preferred_language,
                        "answer": final_answer
                    }
                )

                qa_report = create_qa_report(
                    query=query,
                    preferred_language=preferred_language,
                    final_answer=final_answer,
                    results=results
                )

                qa_docx = create_qa_docx(
                    query=query,
                    preferred_language=preferred_language,
                    final_answer=final_answer,
                    results=results
                )

                st.session_state.latest_qa_report = qa_report
                st.session_state.latest_qa_docx = qa_docx

                st.markdown("### Answer")

                safe_answer = html.escape(final_answer).replace("\n", "<br>")

                st.markdown(
                    f"""
                    <div class="answer-box">
                        {safe_answer}
                    </div>
                    """,
                    unsafe_allow_html=True
                )

                download_col1, download_col2 = st.columns(2)

                with download_col1:
                    st.download_button(
                        label="Download Answer TXT",
                        data=qa_report,
                        file_name="answer_report.txt",
                        mime="text/plain"
                    )

                with download_col2:
                    st.download_button(
                        label="Download Answer DOCX",
                        data=qa_docx,
                        file_name="answer_report.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

                st.markdown("### Evidence")

                for i, result in enumerate(results, start=1):
                    with st.expander(
                        f"Evidence {i} | {result['source']} | Score: {result['score']:.4f}"
                    ):
                        translated_evidence = translate_text(
                            text=result["chunk"],
                            target_language=preferred_language
                        )

                        st.markdown("**Translated Evidence**")
                        st.write(translated_evidence)

                        st.markdown("---")

                        st.markdown("**Original Evidence**")
                        st.write(result["chunk"])


# -------------------------------
# TAB 4: SUMMARIZE
# -------------------------------
with tab4:
    st.markdown('<div class="section-label">Summarize your document</div>', unsafe_allow_html=True)

    st.markdown(
        """
        <div class="mini-note">
        Generate a 4-point summary and key points from the uploaded document in your preferred language.
        </div>
        """,
        unsafe_allow_html=True
    )

    if not st.session_state.formatted_texts:
        st.warning("Please upload and process a document first.")

    else:
        selected_summary_file = st.selectbox(
            "Select document",
            list(st.session_state.formatted_texts.keys()),
            key="summary_selectbox"
        )

        st.write(f"Summary language: **{preferred_language}**")

        if st.button("Generate Summary"):
            document_text = st.session_state.formatted_texts[selected_summary_file]

            with st.spinner("Generating summary..."):
                summary = generate_document_summary(
                    text=document_text,
                    target_language=preferred_language
                )

            with st.spinner("Extracting key points..."):
                key_points = generate_document_key_points(
                    text=document_text,
                    target_language=preferred_language
                )

            st.session_state.document_summary = summary
            st.session_state.document_key_points = key_points
            st.session_state.summary_file = selected_summary_file
            st.session_state.summary_language = preferred_language
            st.session_state.summary_report = create_summary_report(
                file_name=selected_summary_file,
                target_language=preferred_language,
                summary=summary,
                key_points=key_points
            )
            st.session_state.summary_docx = create_summary_docx(
                file_name=selected_summary_file,
                target_language=preferred_language,
                summary=summary,
                key_points=key_points
            )

            st.success("Summary generated.")

        if st.session_state.document_summary and st.session_state.document_key_points:
            st.markdown("### Summary")

            safe_summary = html.escape(
                st.session_state.document_summary
            ).replace("\n", "<br>")

            st.markdown(
                f"""
                <div class="summary-box">
                    {safe_summary}
                </div>
                """,
                unsafe_allow_html=True
            )

            st.markdown("### Key Points")

            safe_points = html.escape(
                st.session_state.document_key_points
            ).replace("\n", "<br>")

            st.markdown(
                f"""
                <div class="answer-box">
                    {safe_points}
                </div>
                """,
                unsafe_allow_html=True
            )

            download_col1, download_col2 = st.columns(2)

            with download_col1:
                st.download_button(
                    label="Download Summary TXT",
                    data=st.session_state.summary_report,
                    file_name="summary_report.txt",
                    mime="text/plain"
                )

            with download_col2:
                st.download_button(
                    label="Download Summary DOCX",
                    data=st.session_state.summary_docx,
                    file_name="summary_report.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )


# -------------------------------
# TAB 5: HISTORY
# -------------------------------
with tab5:
    st.markdown('<div class="section-label">Recent questions</div>', unsafe_allow_html=True)

    if st.session_state.query_history:
        for item in reversed(st.session_state.query_history):
            with st.expander(
                f"{item['time']} | {item['language']} | {item['question']}"
            ):
                st.markdown("**Question**")
                st.write(item["question"])

                st.markdown("**Answer**")
                st.write(item["answer"])

    else:
        st.info("No questions yet.")
