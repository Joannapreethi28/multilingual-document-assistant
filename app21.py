import os
import re
import shutil
import zipfile
from io import BytesIO
from collections import defaultdict, deque
from datetime import datetime
import html
import streamlit as st
import numpy as np
import faiss
import torch
from pypdf import PdfReader
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from sentence_transformers import SentenceTransformer
from transformers import AutoTokenizer, AutoModelForSeq2SeqLM
from deep_translator import GoogleTranslator
# -------------------------------
# PAGE CONFIG
# -------------------------------
st.set_page_config(
    page_title="Multilingual Document Assistant",
    layout="wide"
)
# -------------------------------
# PREMIUM NATIVE APPLICATION CSS
# -------------------------------
st.markdown(
    """
    <div class="top-brand-accent"></div>
    <style>
    /* -------------------------
       SaaS Typography Base
    --------------------------*/
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800&family=JetBrains+Mono:wght@400;500&display=swap');
    html, body, [class*="css"] {
        font-family: 'Inter', sans-serif;
    }
    /* -------------------------
       Application Top Workspace Brand Line
    --------------------------*/
    .top-brand-accent {
        height: 4px;
        width: 100%;
        background: linear-gradient(90deg, #6366F1 0%, #A855F7 50%, #EC4899 100%);
        position: fixed;
        top: 0;
        left: 0;
        z-index: 99999;
    }
    /* -------------------------
       Workspace Window Bounds
    --------------------------*/
    .block-container {
        max-width: 1500px;
        padding-top: 2.2rem;
        padding-bottom: 4rem;
    }
    /* -------------------------
       Action-Oriented Form Buttons
    --------------------------*/
    div.stButton > button {
        width: 100%;
        border: 1px solid rgba(99, 102, 241, 0.2);
        border-radius: 8px;
        background: #6366F1;
        color: #FFFFFF !important;
        padding: 10px 16px;
        font-weight: 500;
        font-size: 14px;
        box-shadow: 0 2px 6px rgba(99, 102, 241, 0.15);
        transition: all 0.2s ease;
    }
    div.stButton > button:hover {
        background: #4F46E5;
        border-color: #6366F1;
        box-shadow: 0 4px 12px rgba(99, 102, 241, 0.3);
    }
    /* -------------------------
       Export/Download Control Elements
    --------------------------*/
    div.stDownloadButton > button {
        width: 100%;
        border-radius: 8px;
        background: rgba(16, 185, 129, 0.08) !important;
        border: 1px solid rgba(16, 185, 129, 0.25) !important;
        color: #059669 !important;
        font-weight: 600;
        padding: 10px 16px;
        transition: all 0.2s ease;
    }
    div.stDownloadButton > button:hover {
        background: rgba(16, 185, 129, 0.15) !important;
        border-color: #10B981 !important;
        box-shadow: 0 4px 12px rgba(16, 185, 129, 0.1);
    }
    /* -------------------------
       Structured Metric Display Boards
    --------------------------*/
    div[data-testid="stMetric"] {
        background: var(--background-color, rgba(128, 128, 128, 0.04)) !important;
        backdrop-filter: blur(8px);
        -webkit-backdrop-filter: blur(8px);
        border-radius: 12px;
        padding: 16px 20px;
        border: 1px solid rgba(128, 128, 128, 0.15);
        box-shadow: 0 4px 12px rgba(0, 0, 0, 0.02);
    }
    div[data-testid="stMetric"] label {
        font-size: 12px !important;
        font-weight: 600 !important;
        text-transform: uppercase;
        letter-spacing: 0.5px;
        opacity: 0.7;
    }
    div[data-testid="stMetric"] div {
        font-size: 24px !important;
        font-weight: 700 !important;
        color: #6366F1 !important;
    }
    /* -------------------------
       Sidebar Branding Matrix
    --------------------------*/
    .sidebar-branding {
        padding: 10px 0 20px 0;
        border-bottom: 1px solid rgba(128, 128, 128, 0.15);
        margin-bottom: 25px;
    }
    .sidebar-title {
        font-size: 21px;
        font-weight: 800;
        letter-spacing: -0.5px;
        margin: 0;
    }
    .sidebar-title span {
        color: #6366F1;
    }
    .sidebar-tag {
        font-size: 11px;
        opacity: 0.5;
        text-transform: uppercase;
        letter-spacing: 0.5px;
    }
    /* -------------------------
       Application Layout Header Bar
    --------------------------*/
    .app-header-container {
        display: flex;
        justify-content: space-between;
        align-items: center;
        background: rgba(128, 128, 128, 0.03);
        border: 1px solid rgba(128, 128, 128, 0.15);
        padding: 20px 28px;
        border-radius: 12px;
        margin-bottom: 25px;
    }
    .app-meta-branding {
        display: flex;
        align-items: center;
        gap: 12px;
    }
    .app-main-title {
        font-size: 22px;
        font-weight: 800;
        letter-spacing: -0.5px;
        margin: 0;
    }
    .app-status-badge-pill {
        background: rgba(99, 102, 241, 0.1);
        color: #6366F1;
        padding: 3px 9px;
        border-radius: 20px;
        font-size: 10px;
        font-weight: 600;
        text-transform: uppercase;
        border: 1px solid rgba(99, 102, 241, 0.25);
    }
    .app-meta-desc {
        font-size: 13px;
        opacity: 0.7;
        margin-top: 4px;
    }
    /* -------------------------
       Functional Workspace Panels
    --------------------------*/
    .panel-header-title {
        font-size: 14px;
        font-weight: 700;
        text-transform: uppercase;
        letter-spacing: 0.5px;
        margin-bottom: 14px;
        opacity: 0.9;
    }
    /* -------------------------
       Analytical Visual Presentation Elements
    --------------------------*/
    .answer-box {
        padding: 20px;
        border-radius: 12px;
        background: rgba(16, 185, 129, 0.04);
        border-left: 5px solid #10B981;
        border-top: 1px solid rgba(128, 128, 128, 0.15);
        border-right: 1px solid rgba(128, 128, 128, 0.15);
        border-bottom: 1px solid rgba(128, 128, 128, 0.15);
        margin-top: 12px;
        margin-bottom: 22px;
        line-height: 1.65;
        font-size: 15px;
        white-space: pre-wrap;
    }
    .summary-box {
        padding: 20px;
        border-radius: 12px;
        background: rgba(245, 158, 11, 0.04);
        border-left: 5px solid #F59E0B;
        border-top: 1px solid rgba(128, 128, 128, 0.15);
        border-right: 1px solid rgba(128, 128, 128, 0.15);
        border-bottom: 1px solid rgba(128, 128, 128, 0.15);
        margin-top: 12px;
        margin-bottom: 22px;
        line-height: 1.65;
        font-size: 15px;
        white-space: pre-wrap;
    }
    .mini-note {
        padding: 14px 18px;
        border-radius: 10px;
        background: rgba(128, 128, 128, 0.03);
        border: 1px solid rgba(128, 128, 128, 0.15);
        margin-bottom: 18px;
        font-size: 13.5px;
        line-height: 1.5;
    }
    .stTextArea textarea {
        border-radius: 12px !important;
        line-height: 1.6;
    }
    section[data-testid="stFileUploader"] {
        padding: 18px;
        border-radius: 12px;
        background-color: transparent !important;
        border: 2px dashed rgba(128, 128, 128, 0.25) !important;
    }
    /* -------------------------
       Workspace Form Bounds
    --------------------------*/
    div[data-baseweb="select"], input, textarea {
        background-color: var(--background-color, #FFFFFF) !important;
        border-radius: 8px !important;
        border: 1px solid rgba(128, 128, 128, 0.25) !important;
    }
    hr {
        border: none;
        border-top: 1px solid rgba(128, 128, 128, 0.15);
        margin: 20px 0;
    }
    .muted {
        opacity: 0.6;
        font-size: 13px;
    }
    </style>
    """,
    unsafe_allow_html=True
)
# -------------------------------
# PREMIUM VISUAL POLISH LAYER
# -------------------------------
st.markdown(
    """
    <style>
    /* Soft, layered application backdrop */
    [data-testid="stAppViewContainer"] {
        background:
            radial-gradient(1100px 560px at 100% -12%, rgba(168, 85, 247, 0.10), transparent 60%),
            radial-gradient(950px 520px at -12% 8%, rgba(99, 102, 241, 0.10), transparent 55%),
            radial-gradient(800px 420px at 60% 120%, rgba(236, 72, 153, 0.06), transparent 60%);
    }

    /* Section headers get a small gradient tab */
    .panel-header-title {
        display: flex;
        align-items: center;
        gap: 10px;
        font-size: 15px !important;
    }
    .panel-header-title::before {
        content: "";
        width: 18px;
        height: 3px;
        border-radius: 3px;
        background: linear-gradient(90deg, #6366F1, #A855F7);
        display: inline-block;
    }

    /* Header bar: gradient wash, accent rail, gradient title */
    .app-header-container {
        background: linear-gradient(135deg, rgba(99,102,241,0.12), rgba(168,85,247,0.07) 55%, rgba(236,72,153,0.05)) !important;
        border: 1px solid rgba(99, 102, 241, 0.18) !important;
        box-shadow: 0 12px 34px rgba(99, 102, 241, 0.10) !important;
        position: relative;
        overflow: hidden;
    }
    .app-header-container::before {
        content: "";
        position: absolute;
        top: 0; bottom: 0; left: 0;
        width: 5px;
        background: linear-gradient(180deg, #6366F1, #A855F7, #EC4899);
    }
    .app-main-title {
        background: linear-gradient(90deg, #6366F1, #A855F7 60%, #EC4899);
        -webkit-background-clip: text;
        background-clip: text;
        -webkit-text-fill-color: transparent;
    }
    .app-status-badge-pill {
        background: linear-gradient(90deg, rgba(99,102,241,0.16), rgba(168,85,247,0.16)) !important;
        box-shadow: inset 0 0 0 1px rgba(99, 102, 241, 0.18);
    }

    /* Metric cards: top accent, depth, hover lift */
    div[data-testid="stMetric"] {
        position: relative;
        overflow: hidden;
        border: 1px solid rgba(99, 102, 241, 0.14) !important;
        box-shadow: 0 8px 24px rgba(15, 23, 42, 0.06) !important;
        transition: transform 0.2s ease, box-shadow 0.2s ease;
    }
    div[data-testid="stMetric"]::before {
        content: "";
        position: absolute;
        top: 0; left: 0; right: 0;
        height: 3px;
        background: linear-gradient(90deg, #6366F1, #A855F7, #EC4899);
    }
    div[data-testid="stMetric"]:hover {
        transform: translateY(-3px);
        box-shadow: 0 16px 34px rgba(99, 102, 241, 0.16) !important;
    }
    div[data-testid="stMetric"] div[data-testid="stMetricValue"] {
        background: linear-gradient(90deg, #6366F1, #A855F7);
        -webkit-background-clip: text;
        background-clip: text;
        -webkit-text-fill-color: transparent;
    }

    /* Primary buttons: gradient fill + hover lift */
    div.stButton > button {
        background: linear-gradient(135deg, #6366F1, #7C3AED) !important;
        border: none !important;
        box-shadow: 0 6px 16px rgba(99, 102, 241, 0.25) !important;
        letter-spacing: 0.2px;
    }
    div.stButton > button:hover {
        background: linear-gradient(135deg, #4F46E5, #6D28D9) !important;
        transform: translateY(-1px);
        box-shadow: 0 10px 24px rgba(99, 102, 241, 0.34) !important;
    }

    /* Download buttons: soft emerald gradient */
    div.stDownloadButton > button {
        background: linear-gradient(135deg, rgba(16,185,129,0.14), rgba(5,150,105,0.10)) !important;
        box-shadow: 0 6px 16px rgba(16, 185, 129, 0.12) !important;
    }
    div.stDownloadButton > button:hover {
        transform: translateY(-1px);
    }

    /* Result cards: subtle depth */
    .answer-box  { box-shadow: 0 10px 26px rgba(16, 185, 129, 0.08); }
    .summary-box { box-shadow: 0 10px 26px rgba(245, 158, 11, 0.08); }
    .mini-note {
        background: linear-gradient(135deg, rgba(99,102,241,0.06), rgba(168,85,247,0.04)) !important;
        border: 1px solid rgba(99, 102, 241, 0.15) !important;
    }

    /* Sidebar shell + brand mark */
    [data-testid="stSidebar"] {
        background: linear-gradient(180deg, rgba(99, 102, 241, 0.06), transparent 32%);
        border-right: 1px solid rgba(128, 128, 128, 0.12);
    }
    .sidebar-branding { border-bottom: 1px solid rgba(99, 102, 241, 0.15); }
    .brand-row { display: flex; align-items: center; gap: 12px; }
    .brand-logo {
        width: 42px; height: 42px;
        border-radius: 12px;
        background: linear-gradient(135deg, #6366F1, #A855F7, #EC4899);
        color: #FFFFFF;
        font-weight: 800;
        font-size: 16px;
        letter-spacing: 0.5px;
        display: flex; align-items: center; justify-content: center;
        box-shadow: 0 8px 18px rgba(99, 102, 241, 0.38);
    }
    .sidebar-title { font-size: 23px; }
    .sidebar-title span {
        background: linear-gradient(90deg, #6366F1, #A855F7, #EC4899);
        -webkit-background-clip: text;
        background-clip: text;
        -webkit-text-fill-color: transparent;
    }

    /* File uploader dropzone: gradient tint + hover */
    section[data-testid="stFileUploader"] {
        transition: all 0.2s ease;
        background: linear-gradient(135deg, rgba(99,102,241,0.04), rgba(168,85,247,0.03)) !important;
    }
    section[data-testid="stFileUploader"]:hover {
        border-color: rgba(99, 102, 241, 0.5) !important;
        background: rgba(99, 102, 241, 0.06) !important;
    }

    /* Inputs: focus ring */
    .stTextInput input:focus,
    textarea:focus,
    div[data-baseweb="select"]:focus-within {
        border-color: #6366F1 !important;
        box-shadow: 0 0 0 3px rgba(99, 102, 241, 0.15) !important;
    }

    /* Expanders: rounded, bordered, subtle shadow */
    div[data-testid="stExpander"] {
        border: 1px solid rgba(99, 102, 241, 0.15) !important;
        border-radius: 12px !important;
        overflow: hidden;
        box-shadow: 0 4px 14px rgba(15, 23, 42, 0.04);
    }

    /* Content headings polish */
    h3 { font-weight: 700 !important; letter-spacing: -0.3px; }
    </style>
    """,
    unsafe_allow_html=True
)
# -------------------------------
# LANGUAGE CODES
# -------------------------------
LANGUAGE_CODES = {
    "English": "en",
    "Tamil": "ta",
    "Hindi": "hi",
    "French": "fr",
    "Spanish": "es"
}
# -------------------------------
# REPORT LABELS
# -------------------------------
REPORT_LABELS = {
    "English": {
        "translation_report": "Document Translation Report",
        "qa_report": "Question Answering Report",
        "summary_report": "Document Summary Report",
        "generated_on": "Generated On",
        "source_file": "Source File",
        "translated_language": "Translated Language",
        "answer_language": "Answer Language",
        "summary_language": "Summary Language",
        "original_document": "Original Document",
        "translated_document": "Translated Document",
        "question": "Question",
        "answer": "Answer",
        "translated_evidence": "Translated Evidence",
        "original_evidence": "Original Evidence",
        "evidence": "Evidence",
        "source": "Source",
        "similarity_score": "Similarity Score",
        "summary": "Summary",
        "key_points": "Key Points"
    },
    "Tamil": {
        "translation_report": "ஆவண மொழிபெயர்ப்பு அறிக்கை",
        "qa_report": "கேள்வி பதில் அறிக்கை",
        "summary_report": "ஆவண சுருக்க அறிக்கை",
        "generated_on": "உருவாக்கப்பட்ட நேரம்",
        "source_file": "மூல கோப்பு",
        "translated_language": "மொழிபெயர்க்கப்பட்ட மொழி",
        "answer_language": "பதில் மொழி",
        "summary_language": "சுருக்க மொழி",
        "original_document": "அசல் ஆவணம்",
        "translated_document": "மொழிபெயர்க்கப்பட்ட ஆவணம்",
        "question": "கேள்வி",
        "answer": "பதில்",
        "translated_evidence": "மொழிபெயர்க்கப்பட்ட ஆதாரம்",
        "original_evidence": "அசல் ஆதாரம்",
        "evidence": "ஆதாரம்",
        "source": "மூலம்",
        "similarity_score": "ஒற்றுமை மதிப்பெண்",
        "summary": "சுருக்கம்",
        "key_points": "முக்கிய குறிப்புகள்"
    },
    "Hindi": {
        "translation_report": "दस्तावेज़ अनुवाद रिपोर्ट",
        "qa_report": "प्रश्न उत्तर रिपोर्ट",
        "summary_report": "दस्तावेज़ सारांश रिपोर्ट",
        "generated_on": "तैयार किया गया",
        "source_file": "स्रोत फ़ाइल",
        "translated_language": "अनुवादित भाषा",
        "answer_language": "उत्तर की भाषा",
        "summary_language": "सारांश की भाषा",
        "original_document": "मूल दस्तावेज़",
        "translated_document": "अनुवादित दस्तावेज़",
        "question": "प्रश्न",
        "answer": "उत्तर",
        "translated_evidence": "अनुवादित प्रमाण",
        "original_evidence": "मूल प्रमाण",
        "evidence": "प्रमाण",
        "source": "स्रोत",
        "similarity_score": "समानता स्कोर",
        "summary": "सारांश",
        "key_points": "मुख्य बिंदु"
    },
    "French": {
        "translation_report": "Rapport de traduction du document",
        "qa_report": "Rapport de questions-réponses",
        "summary_report": "Rapport de résumé du document",
        "generated_on": "Généré le",
        "source_file": "Fichier source",
        "translated_language": "Langue traduite",
        "answer_language": "Langue de réponse",
        "summary_language": "Langue du résumé",
        "original_document": "Document original",
        "translated_document": "Document traduit",
        "question": "Question",
        "answer": "Réponse",
        "translated_evidence": "Preuve traduite",
        "original_evidence": "Preuve originale",
        "evidence": "Preuve",
        "source": "Source",
        "similarity_score": "Score de similarité",
        "summary": "Résumé",
        "key_points": "Points clés"
    },
    "Spanish": {
        "translation_report": "Informe de traducción del documento",
        "qa_report": "Informe de preguntas y respuestas",
        "summary_report": "Informe de resumen del documento",
        "generated_on": "Generado el",
        "source_file": "Archivo fuente",
        "translated_language": "Idioma traducido",
        "answer_language": "Idioma de respuesta",
        "summary_language": "Idioma del resumen",
        "original_document": "Documento original",
        "translated_document": "Documento traducido",
        "question": "Pregunta",
        "answer": "Respuesta",
        "translated_evidence": "Evidencia traducida",
        "original_evidence": "Evidencia original",
        "evidence": "Evidencia",
        "source": "Fuente",
        "similarity_score": "Puntuación de similitud",
        "summary": "Resumen",
        "key_points": "Puntos clave"
    }
}
def label(key, language):
    return REPORT_LABELS.get(language, REPORT_LABELS["English"]).get(key, key)
# -------------------------------
# LOAD MODELS
# -------------------------------
@st.cache_resource
def load_embedding_model():
    return SentenceTransformer(
        "sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2"
    )
@st.cache_resource
def load_answer_model():
    model_name = "google/flan-t5-small"
    tokenizer = AutoTokenizer.from_pretrained(model_name)
    model = AutoModelForSeq2SeqLM.from_pretrained(model_name)
    device = "cuda" if torch.cuda.is_available() else "cpu"
    model = model.to(device)
    return tokenizer, model, device
embedding_model = load_embedding_model()
answer_tokenizer, answer_model, device = load_answer_model()
# -------------------------------
# TEXT CLEANING
# -------------------------------
def clean_text(text):
    text = text.replace("\n", " ")
    text = re.sub(r"\s+", " ", text)
    return text.strip()
def split_into_sentences(text):
    sentences = re.split(r"(?<=[.!?।])\s+", text)
    return [sentence.strip() for sentence in sentences if sentence.strip()]
def clean_raw_text(text):
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    return text.strip()
def make_readable_paragraphs(text):
    text = clean_raw_text(text)
    text = re.sub(r"(\d+)\.([A-Za-z])", r"\1. \2", text)
    text = re.sub(r"\s+(\d+\.\s*)", r"\n\n\1", text)
    headings = [
        "Problem Statement", "Abstract", "Proposed Solution", "Tech Stack",
        "Workflow", "System Architecture", "Output", "Conclusion", "Features",
        "Advantages", "Limitations", "Future Enhancement", "Demo Scenario",
        "Modules", "Objective"
    ]
    for heading in headings:
        text = re.sub(
            rf"\s*({heading}\s*:?)",
            rf"\n\n\1\n",
            text,
            flags=re.IGNORECASE
        )
    blocks = re.split(r"\n\s*\n", text)
    final_blocks = []
    for block in blocks:
        block = re.sub(r"\s+", " ", block).strip()
        if not block:
            continue
        if len(block.split()) <= 8 and not block.endswith("."):
            final_blocks.append(block)
            continue
        sentences = split_into_sentences(block)
        if not sentences:
            final_blocks.append(block)
            continue
        current_para = ""
        for sentence in sentences:
            if len(current_para) + len(sentence) <= 700:
                if current_para:
                    current_para += " " + sentence
                else:
                    current_para = sentence
            else:
                if current_para:
                    final_blocks.append(current_para.strip())
                current_para = sentence
        if current_para:
            final_blocks.append(current_para.strip())
    return "\n\n".join(final_blocks).strip()
def fix_broken_pdf_lines(text):
    text = clean_raw_text(text)
    lines = [line.strip() for line in text.split("\n") if line.strip()]
    if not lines:
        return ""
    small_line_count = 0
    for line in lines:
        if len(line.split()) <= 2:
            small_line_count += 1
    small_line_ratio = small_line_count / len(lines)
    if small_line_ratio > 0.45:
        merged_text = " ".join(lines)
        merged_text = re.sub(r"\s+", " ", merged_text).strip()
        return make_readable_paragraphs(merged_text)
    paragraphs = []
    current_para = ""
    for raw_line in text.split("\n"):
        line = raw_line.strip()
        if not line:
            if current_para.strip():
                paragraphs.append(current_para.strip())
                current_para = ""
            continue
        if not current_para:
            current_para = line
        else:
            if current_para.endswith("-"):
                current_para = current_para[:-1] + line
            else:
                current_para += " " + line
    if current_para.strip():
        paragraphs.append(current_para.strip())
    return make_readable_paragraphs("\n\n".join(paragraphs))
def clean_for_docx(text):
    text = str(text)
    text = re.sub(r"[\x00-\x08\x0B\x0C\x0E-\x1F]", "", text)
    text = text.replace("\u2028", "\n").replace("\u2029", "\n")
    return text
# -------------------------------
# TEXT EXTRACTION BYTES
# -------------------------------
def extract_pdf_text_for_qa(file_bytes):
    text = ""
    pdf_reader = PdfReader(BytesIO(file_bytes))
    for page in pdf_reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + " "
    return clean_text(text)
def extract_docx_text_for_qa(file_bytes):
    text = ""
    document = Document(BytesIO(file_bytes))
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            text += paragraph.text + " "
    return clean_text(text)
def extract_txt_text_for_qa(file_bytes):
    text = file_bytes.decode("utf-8", errors="ignore")
    return clean_text(text)
def extract_text_for_qa(file_name, file_bytes):
    file_name = file_name.lower()
    if file_name.endswith(".pdf"):
        return extract_pdf_text_for_qa(file_bytes)
    elif file_name.endswith(".docx"):
        return extract_docx_text_for_qa(file_bytes)
    elif file_name.endswith(".txt"):
        return extract_txt_text_for_qa(file_bytes)
    return ""
def extract_pdf_text_for_translation(file_bytes):
    page_texts = []
    pdf_reader = PdfReader(BytesIO(file_bytes))
    for page in pdf_reader.pages:
        page_text = page.extract_text()
        if page_text:
            page_texts.append(page_text)
    return fix_broken_pdf_lines("\n\n".join(page_texts))
def extract_docx_text_for_translation(file_bytes):
    document = Document(BytesIO(file_bytes))
    paragraphs = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            paragraphs.append(text)
    return make_readable_paragraphs("\n\n".join(paragraphs))
def extract_txt_text_for_translation(file_bytes):
    raw_text = file_bytes.decode("utf-8", errors="ignore")
    return make_readable_paragraphs(raw_text)
def extract_text_for_translation(file_name, file_bytes):
    file_name = file_name.lower()
    if file_name.endswith(".pdf"):
        return extract_pdf_text_for_translation(file_bytes)
    elif file_name.endswith(".docx"):
        return extract_docx_text_for_translation(file_bytes)
    elif file_name.endswith(".txt"):
        return extract_txt_text_for_translation(file_bytes)
    return ""
# -------------------------------
# CHUNKING
# -------------------------------
def chunk_text(text, chunk_size=160, overlap=35):
    words = text.split()
    chunks = []
    start = 0
    while start < len(words):
        end = start + chunk_size
        chunk = " ".join(words[start:end])
        if chunk.strip():
            chunks.append(chunk.strip())
        start = end - overlap
    return chunks
def split_long_text_for_translation(text, max_chars=2800):
    if len(text) <= max_chars:
        return [text]
    sentences = split_into_sentences(text)
    if not sentences:
        return [text[i:i + max_chars] for i in range(0, len(text), max_chars)]
    parts = []
    current_part = ""
    for sentence in sentences:
        if len(sentence) > max_chars:
            if current_part.strip():
                parts.append(current_part.strip())
                current_part = ""
            for i in range(0, len(sentence), max_chars):
                parts.append(sentence[i:i + max_chars])
        elif len(current_part) + len(sentence) + 1 <= max_chars:
            if current_part:
                current_part += " " + sentence
            else:
                current_part = sentence
        else:
            if current_part.strip():
                parts.append(current_part.strip())
            current_part = sentence
    if current_part.strip():
        parts.append(current_part.strip())
    return parts
# -------------------------------
# TRANSLATION
# -------------------------------
def translate_to_english(text):
    try:
        return GoogleTranslator(source="auto", target="en").translate(text)
    except Exception:
        return text
def translate_text(text, target_language):
    target_code = LANGUAGE_CODES.get(target_language, "en")
    if target_code == "en":
        return text
    try:
        return GoogleTranslator(source="auto", target=target_code).translate(text)
    except Exception:
        return text
def translate_full_document(text, target_language):
    target_code = LANGUAGE_CODES.get(target_language, "en")
    if target_code == "en":
        return text
    paragraphs = [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]
    translation_units = []
    for paragraph in paragraphs:
        parts = split_long_text_for_translation(paragraph)
        translation_units.append(parts)
    total_parts = sum(len(parts) for parts in translation_units)
    if total_parts == 0:
        return ""
    progress_bar = st.progress(0)
    status_text = st.empty()
    completed_parts = 0
    translated_paragraphs = []
    for parts in translation_units:
        translated_parts = []
        for part in parts:
            status_text.write(f"Translating subsection {completed_parts + 1} of {total_parts}...")
            try:
                translated_part = GoogleTranslator(source="auto", target=target_code).translate(part)
            except Exception:
                translated_part = part
            translated_parts.append(translated_part)
            completed_parts += 1
            progress_bar.progress(completed_parts / total_parts)
        translated_paragraphs.append(" ".join(translated_parts).strip())
    status_text.empty()
    return "\n\n".join(translated_paragraphs).strip()
# -------------------------------
# FAISS SEARCH
# -------------------------------
def create_faiss_index(chunks):
    embeddings = embedding_model.encode(chunks, convert_to_numpy=True)
    embeddings = np.array(embeddings).astype("float32")
    faiss.normalize_L2(embeddings)
    dimension = embeddings.shape[1]
    index = faiss.IndexFlatIP(dimension)
    index.add(embeddings)
    return index
def search_query(query, index, chunks, chunk_sources, top_k=3):
    query_embedding = embedding_model.encode([query], convert_to_numpy=True)
    query_embedding = np.array(query_embedding).astype("float32")
    faiss.normalize_L2(query_embedding)
    scores, indices = index.search(query_embedding, top_k)
    results = []
    for score, idx in zip(scores[0], indices[0]):
        if idx != -1:
            results.append({
                "chunk": chunks[idx],
                "source": chunk_sources[idx],
                "score": float(score)
            })
    return results
# -------------------------------
# SLM GENERATION GENERATION
# -------------------------------
def generate_with_slm(prompt, max_new_tokens=180):
    inputs = answer_tokenizer(prompt, return_tensors="pt", truncation=True, max_length=1024).to(device)
    with torch.no_grad():
        outputs = answer_model.generate(
            **inputs,
            max_new_tokens=max_new_tokens,
            num_beams=4,
            no_repeat_ngram_size=3,
            early_stopping=True
        )
    return answer_tokenizer.decode(outputs[0], skip_special_tokens=True).strip()
def extractive_answer(query, chunks, max_sentences=4):
    all_sentences = []
    for chunk in chunks:
        all_sentences.extend(split_into_sentences(chunk))
    if not all_sentences:
        return "The answer is not available in the uploaded documents."
    sentence_embeddings = embedding_model.encode(all_sentences, convert_to_numpy=True)
    query_embedding = embedding_model.encode([query], convert_to_numpy=True)
    sentence_embeddings = np.array(sentence_embeddings).astype("float32")
    query_embedding = np.array(query_embedding).astype("float32")
    faiss.normalize_L2(sentence_embeddings)
    faiss.normalize_L2(query_embedding)
    scores = np.dot(sentence_embeddings, query_embedding.T).flatten()
    top_indices = scores.argsort()[-max_sentences:][::-1]
    selected_sentences = []
    for idx in top_indices:
        sentence = all_sentences[idx]
        if sentence not in selected_sentences:
            selected_sentences.append(sentence)
    return " ".join(selected_sentences)
def generate_clean_answer(query, retrieved_chunks):
    english_query = translate_to_english(query)
    english_context_parts = []
    for chunk in retrieved_chunks[:3]:
        english_context_parts.append(translate_to_english(chunk[:3500]))
    context = "\n\n".join(english_context_parts)
    prompt = f"""
Answer the question using only the given context.
Rules:
- Give a clear and simple answer.
- Do not copy the full context.
- Do not add outside information.
- If the answer is not found, say: The answer is not available in the uploaded documents.
Context:
{context}
Question:
{english_query}
Answer:
"""
    answer = generate_with_slm(prompt, max_new_tokens=180)
    weak_answers = ["", "answer:", "context:", "The answer is not available in the uploaded documents."]
    if answer.lower() in [item.lower() for item in weak_answers] or len(answer.split()) < 8:
        answer = extractive_answer(query=english_query, chunks=english_context_parts)
    return answer
# -------------------------------
# SUMMARY ENGINES
# -------------------------------
def prepare_summary_text(text, max_chars=4500):
    paragraphs = [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]
    selected_text = ""
    for paragraph in paragraphs:
        if len(selected_text) + len(paragraph) + 2 <= max_chars:
            selected_text += paragraph + "\n\n"
        else:
            break
    return selected_text.strip() if selected_text.strip() else text[:max_chars]
def fallback_summary(text):
    sentences = split_into_sentences(text)[:4]
    if not sentences:
        return "1. Summary is not available for this document."
    return "\n".join(f"{i}. {sentence}" for i, sentence in enumerate(sentences, start=1))
def fallback_key_points(text):
    sentences = split_into_sentences(text)[:5]
    if not sentences:
        return "- Key points are not available for this document."
    return "\n".join(f"- {sentence}" for sentence in sentences)
def generate_document_summary(text, target_language):
    source_text = prepare_summary_text(text, max_chars=4500)
    english_text = translate_to_english(source_text)
    prompt = f"""
Read the following document and create a clear summary.
Rules:
- Write the summary as 4 numbered points.
- Each point should be one full sentence.
- Keep it simple and useful.
- Do not add outside information.
Document:
{english_text}
4-point summary:
"""
    summary = generate_with_slm(prompt, max_new_tokens=260)
    if not summary or len(summary.split()) < 35:
        summary = fallback_summary(english_text)
    if "1." not in summary and "- " not in summary:
        sentences = split_into_sentences(summary)
        if len(sentences) >= 3:
            summary = "\n".join(f"{i}. {sentence}" for i, sentence in enumerate(sentences[:4], start=1))
        else:
            summary = fallback_summary(english_text)
    return translate_text(text=summary, target_language=target_language)
def generate_document_key_points(text, target_language):
    source_text = prepare_summary_text(text, max_chars=4500)
    english_text = translate_to_english(source_text)
    prompt = f"""
Extract 5 important key points from the following document.
Write each point as a short bullet.
Do not add outside information.
Document:
{english_text}
Key points:
"""
    key_points = generate_with_slm(prompt, max_new_tokens=220)
    if not key_points or len(key_points.split()) < 10:
        key_points = fallback_key_points(english_text)
    return translate_text(text=key_points, target_language=target_language)
# -------------------------------
# DOCX COMPILERS
# -------------------------------
def set_run_font(run, size=11, bold=False):
    run.font.name = "Nirmala UI"
    run.font.size = Pt(size)
    run.font.bold = bold
    run_element = run._element
    run_properties = run_element.get_or_add_rPr()
    run_fonts = run_properties.get_or_add_rFonts()
    run_fonts.set(qn("w:ascii"), "Nirmala UI")
    run_fonts.set(qn("w:hAnsi"), "Nirmala UI")
    run_fonts.set(qn("w:eastAsia"), "Nirmala UI")
    run_fonts.set(qn("w:cs"), "Nirmala UI")
def split_long_line(line, limit=1800):
    return [line[i:i + limit] for i in range(0, len(line), limit)] if len(line) > limit else [line]
def add_multiline_text(document, text):
    text = clean_for_docx(text)
    for line in text.split("\n"):
        if not line.strip():
            document.add_paragraph("")
            continue
        for part in split_long_line(line.strip()):
            paragraph = document.add_paragraph()
            run = paragraph.add_run(part)
            set_run_font(run, size=11)
def create_docx_file(title, sections):
    document = Document()
    styles = document.styles
    styles["Normal"].font.name = "Nirmala UI"
    styles["Normal"].font.size = Pt(11)
    heading = document.add_heading("", level=0)
    heading_run = heading.add_run(clean_for_docx(title))
    set_run_font(heading_run, size=18, bold=True)
    generated_para = document.add_paragraph()
    generated_run = generated_para.add_run(f"Generated On: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    set_run_font(generated_run, size=10)
    for section_title, section_text in sections:
        section_heading = document.add_heading("", level=1)
        section_run = section_heading.add_run(clean_for_docx(section_title))
        set_run_font(section_run, size=14, bold=True)
        add_multiline_text(document, section_text)
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
# -------------------------------
# SYSTEM REPORT OUTPUT GENERATION
# -------------------------------
def create_qa_report(query, preferred_language, final_answer, results):
    report = label("qa_report", preferred_language).upper() + "\n" + "=" * 60 + "\n\n"
    report += f"{label('generated_on', preferred_language)}: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
    report += f"{label('answer_language', preferred_language)}: {preferred_language}\n\n"
    report += label("question", preferred_language).upper() + "\n" + "-" * 60 + "\n" + query + "\n\n"
    report += label("answer", preferred_language).upper() + "\n" + "-" * 60 + "\n" + final_answer + "\n\n"
    report += label("evidence", preferred_language).upper() + "\n" + "-" * 60 + "\n"
    for i, result in enumerate(results, start=1):
        translated_evidence = translate_text(text=result["chunk"], target_language=preferred_language)
        report += f"\n{label('evidence', preferred_language)} {i}\n"
        report += f"{label('source', preferred_language)}: {result['source']}\n"
        report += f"{label('similarity_score', preferred_language)}: {result['score']:.4f}\n\n"
        report += label("translated_evidence", preferred_language) + "\n" + translated_evidence + "\n\n"
        report += label("original_evidence", preferred_language) + "\n" + result["chunk"] + "\n"
    return report
def create_qa_docx(query, preferred_language, final_answer, results):
    evidence_text = ""
    for i, result in enumerate(results, start=1):
        translated_evidence = translate_text(text=result["chunk"], target_language=preferred_language)
        evidence_text += f"{label('evidence', preferred_language)} {i}\n"
        evidence_text += f"{label('source', preferred_language)}: {result['source']}\n"
        evidence_text += f"{label('similarity_score', preferred_language)}: {result['score']:.4f}\n\n"
        evidence_text += f"{label('translated_evidence', preferred_language)}:\n{translated_evidence}\n\n"
        evidence_text += f"{label('original_evidence', preferred_language)}:\n{result['chunk']}\n\n"
    sections = [
        (label("question", preferred_language), query),
        (label("answer", preferred_language), final_answer),
        (label("evidence", preferred_language), evidence_text)
    ]
    return create_docx_file(title=label("qa_report", preferred_language), sections=sections)
def create_translation_report(file_name, target_language, original_text, translated_text):
    report = label("translation_report", target_language).upper() + "\n" + "=" * 60 + "\n\n"
    report += f"{label('generated_on', target_language)}: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
    report += f"{label('source_file', target_language)}: {file_name}\n"
    report += f"{label('translated_language', target_language)}: {target_language}\n\n"
    report += label("original_document", target_language).upper() + "\n" + "-" * 60 + "\n" + original_text + "\n\n"
    report += label("translated_document", target_language).upper() + "\n" + "-" * 60 + "\n\n" + translated_text
    return report
def create_translation_docx(file_name, target_language, original_text, translated_text):
    sections = [
        (label("source_file", target_language), file_name),
        (label("original_document", target_language), original_text),
        (label("translated_document", target_language), translated_text)
    ]
    return create_docx_file(title=label("translation_report", target_language), sections=sections)
def create_summary_report(file_name, target_language, summary, key_points):
    report = label("summary_report", target_language).upper() + "\n" + "=" * 60 + "\n\n"
    report += f"{label('generated_on', target_language)}: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
    report += f"{label('source_file', target_language)}: {file_name}\n"
    report += f"{label('summary_language', target_language)}: {target_language}\n\n"
    report += label("summary", target_language).upper() + "\n" + "-" * 60 + "\n\n" + summary + "\n\n"
    report += label("key_points", target_language).upper() + "\n" + "-" * 60 + "\n\n" + key_points
    return report
def create_summary_docx(file_name, target_language, summary, key_points):
    sections = [
        (label("source_file", target_language), file_name),
        (label("summary", target_language), summary),
        (label("key_points", target_language), key_points)
    ]
    return create_docx_file(title=label("summary_report", target_language), sections=sections)
# -------------------------------
# WORKSPACE PERSISTED STATES
# -------------------------------
if "uploaded_docs" not in st.session_state: st.session_state.uploaded_docs = {}
if "formatted_texts" not in st.session_state: st.session_state.formatted_texts = {}
if "chunks" not in st.session_state: st.session_state.chunks = []
if "chunk_sources" not in st.session_state: st.session_state.chunk_sources = []
if "index" not in st.session_state: st.session_state.index = None
if "file_names" not in st.session_state: st.session_state.file_names = []
if "query_history" not in st.session_state: st.session_state.query_history = []
if "latest_qa_report" not in st.session_state: st.session_state.latest_qa_report = None
if "latest_qa_docx" not in st.session_state: st.session_state.latest_qa_docx = None
if "translated_document" not in st.session_state: st.session_state.translated_document = None
if "translation_report" not in st.session_state: st.session_state.translation_report = None
if "translation_docx" not in st.session_state: st.session_state.translation_docx = None
if "selected_translation_file" not in st.session_state: st.session_state.selected_translation_file = None
if "translation_language" not in st.session_state: st.session_state.translation_language = None
if "summary_report" not in st.session_state: st.session_state.summary_report = None
if "summary_docx" not in st.session_state: st.session_state.summary_docx = None
if "document_summary" not in st.session_state: st.session_state.document_summary = None
if "document_key_points" not in st.session_state: st.session_state.document_key_points = None
if "summary_file" not in st.session_state: st.session_state.summary_file = None
if "summary_language" not in st.session_state: st.session_state.summary_language = None
# Operational status trace variables
if "processing_status" not in st.session_state:
    st.session_state.processing_status = "Waiting"
# Deferred translation trigger flag (drives the In Progress / Ready status)
if "pending_translation_file" not in st.session_state:
    st.session_state.pending_translation_file = None
# -------------------------------
# CORE SIDEBAR WORKSPACE ENGINE
# -------------------------------
if "active_view" not in st.session_state:
    st.session_state.active_view = "Upload"
with st.sidebar:
    st.markdown(
        """
        <div class="sidebar-branding">
            <div class="brand-row">
                <div class="brand-logo">DM</div>
                <div>
                    <h1 class="sidebar-title">Docu<span>Mind</span></h1>
                    <div class="sidebar-tag">Intelligence Workspace</div>
                </div>
            </div>
        </div>
        """,
        unsafe_allow_html=True
    )
    is_upload = "nav-item-active" if st.session_state.active_view == "Upload" else ""
    is_translate = "nav-item-active" if st.session_state.active_view == "Translate Document" else ""
    is_qa = "nav-item-active" if st.session_state.active_view == "Ask Questions" else ""
    is_summary = "nav-item-active" if st.session_state.active_view == "Summarize" else ""
    is_history = "nav-item-active" if st.session_state.active_view == "History" else ""
    if st.button("Upload", key="nav_upload", use_container_width=True):
        st.session_state.active_view = "Upload"
        st.rerun()
    if st.button("Translate Document", key="nav_translate", use_container_width=True):
        st.session_state.active_view = "Translate Document"
        st.rerun()
    if st.button("Ask Questions", key="nav_qa", use_container_width=True):
        st.session_state.active_view = "Ask Questions"
        st.rerun()
    if st.button("Summarize", key="nav_summary", use_container_width=True):
        st.session_state.active_view = "Summarize"
        st.rerun()
    if st.button(" History", key="nav_history", use_container_width=True):
        st.session_state.active_view = "History"
        st.rerun()
    st.markdown("<hr>", unsafe_allow_html=True)
    st.markdown('<div class="panel-header-title">Settings</div>', unsafe_allow_html=True)
    preferred_language = st.selectbox(
        "Preferred Language",
        ["English", "Tamil", "Hindi", "French", "Spanish"]
    )
    st.markdown("<br>", unsafe_allow_html=True)
    if st.button("↺ Reset"):
        st.session_state.uploaded_docs = {}
        st.session_state.formatted_texts = {}
        st.session_state.chunks = []
        st.session_state.chunk_sources = []
        st.session_state.index = None
        st.session_state.file_names = []
        st.session_state.query_history = []
        st.session_state.latest_qa_report = None
        st.session_state.latest_qa_docx = None
        st.session_state.translated_document = None
        st.session_state.translation_report = None
        st.session_state.translation_docx = None
        st.session_state.selected_translation_file = None
        st.session_state.translation_language = None
        st.session_state.summary_report = None
        st.session_state.summary_docx = None
        st.session_state.document_summary = None
        st.session_state.document_key_points = None
        st.session_state.summary_file = None
        st.session_state.summary_language = None
        st.session_state.processing_status = "Waiting"
        st.session_state.pending_translation_file = None
        st.success("Reset completed.")
        st.rerun()
    st.markdown(f"""
    <style>
        div[data-testid="stSidebar"] div.stButton:nth-of-type(1) button {{
            background: {"rgba(99, 102, 241, 0.08)" if is_upload else "transparent"} !important;
            color: {"#6366F1" if is_upload else "var(--text-color)"} !important;
            font-weight: {"600" if is_upload else "500"} !important;
            border-left: {"4px solid #6366F1" if is_upload else "1px solid transparent"} !important;
            border-radius: 4px 8px 8px 4px !important;
            text-align: left !important;
            justify-content: flex-start !important;
        }}
        div[data-testid="stSidebar"] div.stButton:nth-of-type(2) button {{
            background: {"rgba(99, 102, 241, 0.08)" if is_translate else "transparent"} !important;
            color: {"#6366F1" if is_translate else "var(--text-color)"} !important;
            font-weight: {"600" if is_translate else "500"} !important;
            border-left: {"4px solid #6366F1" if is_translate else "1px solid transparent"} !important;
            border-radius: 4px 8px 8px 4px !important;
            text-align: left !important;
            justify-content: flex-start !important;
        }}
        div[data-testid="stSidebar"] div.stButton:nth-of-type(3) button {{
            background: {"rgba(99, 102, 241, 0.08)" if is_qa else "transparent"} !important;
            color: {"#6366F1" if is_qa else "var(--text-color)"} !important;
            font-weight: {"600" if is_qa else "500"} !important;
            border-left: {"4px solid #6366F1" if is_qa else "1px solid transparent"} !important;
            border-radius: 4px 8px 8px 4px !important;
            text-align: left !important;
            justify-content: flex-start !important;
        }}
        div[data-testid="stSidebar"] div.stButton:nth-of-type(4) button {{
            background: {"rgba(99, 102, 241, 0.08)" if is_summary else "transparent"} !important;
            color: {"#6366F1" if is_summary else "var(--text-color)"} !important;
            font-weight: {"600" if is_summary else "500"} !important;
            border-left: {"4px solid #6366F1" if is_summary else "1px solid transparent"} !important;
            border-radius: 4px 8px 8px 4px !important;
            text-align: left !important;
            justify-content: flex-start !important;
        }}
        div[data-testid="stSidebar"] div.stButton:nth-of-type(5) button {{
            background: {"rgba(99, 102, 241, 0.08)" if is_history else "transparent"} !important;
            color: {"#6366F1" if is_history else "var(--text-color)"} !important;
            font-weight: {"600" if is_history else "500"} !important;
            border-left: {"4px solid #6366F1" if is_history else "1px solid transparent"} !important;
            border-radius: 4px 8px 8px 4px !important;
            text-align: left !important;
            justify-content: flex-start !important;
        }}
    </style>
    """, unsafe_allow_html=True)
# -------------------------------
# WORKSPACE METRIC BOARD HEADER
# -------------------------------
st.markdown(
    f"""
    <div class="app-header-container">
        <div>
            <div class="app-meta-branding">
                <h2 class="app-main-title">{st.session_state.active_view}</h2>
                <span class="app-status-badge-pill">Active Session</span>
            </div>
            <div class="app-meta-desc">Upload your document once, translate it side-by-side, summarize it, or ask questions from it in your preferred language.</div>
        </div>
    </div>
    """,
    unsafe_allow_html=True,
)
# Operational status dashboard ribbon
db1, db2, db3 = st.columns(3)
with db1: st.metric("Documents", len(st.session_state.file_names))
with db2: st.metric("Text Sections", len(st.session_state.chunks))
with db3: st.metric("Status", st.session_state.processing_status)
# -------------------------------
# VIEW CONTROLLER ROUTER
# -------------------------------
if st.session_state.active_view == "Upload":
    st.markdown('<div class="panel-header-title">Upload your documents</div>', unsafe_allow_html=True)
    st.markdown('<div class="mini-note">Upload once. The same document will be used for translation, summary, and question answering.</div>', unsafe_allow_html=True)
    uploaded_files = st.file_uploader(
        "Choose PDF, TXT, or DOCX files",
        type=["pdf", "txt", "docx"],
        accept_multiple_files=True,
        key="single_upload",
        label_visibility="collapsed"
    )
    if uploaded_files:
        st.write("Selected files:")
        for u_file in uploaded_files:
            st.write(f"- {u_file.name}")
        if st.button("Process Documents"):
            st.session_state.processing_status = "In Progress"
            st.rerun()
    # Deferred status tracking engine iteration wrapper block
    if st.session_state.processing_status == "In Progress" and uploaded_files:
        uploaded_docs, formatted_texts, all_chunks, chunk_sources, file_names = {}, {}, [], [], []
        for uploaded_file in uploaded_files:
            f_name = uploaded_file.name
            f_bytes = uploaded_file.getvalue()
            uploaded_docs[f_name] = f_bytes
            qa_text = extract_text_for_qa(f_name, f_bytes)
            formatted_text = extract_text_for_translation(f_name, f_bytes)
            if formatted_text.strip(): formatted_texts[f_name] = formatted_text
            if qa_text.strip():
                chunks = chunk_text(qa_text)
                for chunk in chunks:
                    all_chunks.append(chunk)
                    chunk_sources.append(f_name)
                file_names.append(f_name)
        if all_chunks:
            st.session_state.uploaded_docs = uploaded_docs
            st.session_state.formatted_texts = formatted_texts
            st.session_state.chunks = all_chunks
            st.session_state.chunk_sources = chunk_sources
            st.session_state.index = create_faiss_index(all_chunks)
            st.session_state.file_names = file_names
            # Drop old logs
            st.session_state.translated_document = None
            st.session_state.translation_report = None
            st.session_state.translation_docx = None
            st.session_state.document_summary = None
            st.session_state.document_key_points = None
            st.session_state.latest_qa_report = None
            st.session_state.latest_qa_docx = None
            st.session_state.processing_status = "Ready"
            st.success("Documents are ready.")
            st.rerun()
        else:
            st.session_state.processing_status = "Waiting"
            st.error("No readable text found in the uploaded documents.")
            st.rerun()
    if st.session_state.file_names:
        st.markdown("<hr>", unsafe_allow_html=True)
        st.markdown("### Ready documents")
        for name in st.session_state.file_names:
            st.write(f"- {name}")
elif st.session_state.active_view == "Translate Document":
    st.markdown('<div class="panel-header-title">Side-by-side document translation</div>', unsafe_allow_html=True)
    if not st.session_state.formatted_texts:
        st.warning("Please upload and process a document first.")
    else:
        selected_file = st.selectbox("Select document", list(st.session_state.formatted_texts.keys()), key="translation_selectbox")
        st.write(f"Translation language: **{preferred_language}**")
        if st.button("Translate Document"):
            # Flag the job and flip the status to In Progress, then rerun so the
            # Status box repaints as "In Progress" BEFORE the translation runs.
            st.session_state.pending_translation_file = selected_file
            st.session_state.processing_status = "In Progress"
            st.rerun()
        # Deferred translation worker: runs on the rerun where the Status box
        # already shows "In Progress", then flips to "Ready" once the result exists.
        if st.session_state.pending_translation_file:
            file_to_translate = st.session_state.pending_translation_file
            orig_text = st.session_state.formatted_texts[file_to_translate]
            with st.spinner("Translating document..."):
                trans_text = translate_full_document(text=orig_text, target_language=preferred_language)
            st.session_state.translated_document = trans_text
            st.session_state.selected_translation_file = file_to_translate
            st.session_state.translation_language = preferred_language
            st.session_state.translation_report = create_translation_report(file_to_translate, preferred_language, orig_text, trans_text)
            st.session_state.translation_docx = create_translation_docx(file_to_translate, preferred_language, orig_text, trans_text)
            st.session_state.pending_translation_file = None
            st.session_state.processing_status = "Ready"
            st.rerun()
        if st.session_state.translated_document:
            st.markdown("### Document View")
            left_canvas, right_canvas = st.columns(2)
            with left_canvas:
                st.markdown("**Original Document**")
                st.text_area("Original", value=st.session_state.formatted_texts[st.session_state.selected_translation_file][:7000], height=500, label_visibility="collapsed")
            with right_canvas:
                st.markdown(f"**Translated Document ({st.session_state.translation_language})**")
                st.text_area("Translated", value=st.session_state.translated_document[:7000], height=500, label_visibility="collapsed")
            if len(st.session_state.translated_document) > 7000:
                st.info("Preview shows the first 7000 characters. Download the complete translation below.")
            down_l, down_r = st.columns(2)
            with down_l: st.download_button("Download Translation TXT", data=st.session_state.translation_report, file_name="translated_document_report.txt", mime="text/plain")
            with down_r: st.download_button("Download Translation DOCX", data=st.session_state.translation_docx, file_name="translated_document_report.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
elif st.session_state.active_view == "Ask Questions":
    st.markdown('<div class="panel-header-title">Ask questions from your document</div>', unsafe_allow_html=True)
    if not st.session_state.chunks or st.session_state.index is None:
        st.warning("Please upload and process a document first.")
    else:
        query_input = st.text_input("Ask your question", placeholder="Example: What is this document about?")
        if st.button("Get Answer"):
            if not query_input.strip():
                st.error("Please enter a question.")
            else:
                with st.spinner("Searching document..."):
                    results = search_query(query=query_input, index=st.session_state.index, chunks=st.session_state.chunks, chunk_sources=st.session_state.chunk_sources, top_k=3)
                
                retrieved_chunks = [res["chunk"] for res in results]
                with st.spinner("Writing answer..."):
                    eng_ans = generate_clean_answer(query=query_input, retrieved_chunks=retrieved_chunks)
                    final_ans = translate_text(text=eng_ans, target_language=preferred_language)
                st.session_state.query_history.append({
                    "time": datetime.now().strftime("%H:%M:%S"),
                    "question": query_input,
                    "language": preferred_language,
                    "answer": final_ans
                })
                st.session_state.latest_qa_report = create_qa_report(query_input, preferred_language, final_ans, results)
                st.session_state.latest_qa_docx = create_qa_docx(query_input, preferred_language, final_ans, results)
                st.markdown("### Answer")
                st.markdown(f'<div class="answer-box">{html.escape(final_ans).replace("\n", "<br>")}</div>', unsafe_allow_html=True)
                down_l, down_r = st.columns(2)
                with down_l: st.download_button("Download Answer TXT", data=st.session_state.latest_qa_report, file_name="answer_report.txt", mime="text/plain")
                with down_r: st.download_button("Download Answer DOCX", data=st.session_state.latest_qa_docx, file_name="answer_report.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                st.markdown("### Evidence")
                for idx, res in enumerate(results, start=1):
                    with st.expander(f"Evidence {idx} | {res['source']} | Score: {res['score']:.4f}"):
                        st.markdown("**Translated Evidence**")
                        st.write(translate_text(res["chunk"], preferred_language))
                        st.markdown("<hr>", unsafe_allow_html=True)
                        st.markdown("**Original Evidence**")
                        st.write(res["chunk"])
elif st.session_state.active_view == "Summarize":
    st.markdown('<div class="panel-header-title">Summarize your document</div>', unsafe_allow_html=True)
    st.markdown('<div class="mini-note">Generate a 4-point summary and key points from the uploaded document in your preferred language.</div>', unsafe_allow_html=True)
    if not st.session_state.formatted_texts:
        st.warning("Please upload and process a document first.")
    else:
        selected_summary_file = st.selectbox("Select document", list(st.session_state.formatted_texts.keys()), key="summary_selectbox")
        st.write(f"Summary language: **{preferred_language}**")
        if st.button("Generate Summary"):
            doc_text = st.session_state.formatted_texts[selected_summary_file]
            with st.spinner("Generating summary..."):
                summary = generate_document_summary(doc_text, preferred_language)
            with st.spinner("Extracting key points..."):
                k_points = generate_document_key_points(doc_text, preferred_language)
            st.session_state.document_summary = summary
            st.session_state.document_key_points = k_points
            st.session_state.summary_file = selected_summary_file
            st.session_state.summary_language = preferred_language
            st.session_state.summary_report = create_summary_report(selected_summary_file, preferred_language, summary, k_points)
            st.session_state.summary_docx = create_summary_docx(selected_summary_file, preferred_language, summary, k_points)
            st.success("Summary generated.")
        if st.session_state.document_summary and st.session_state.document_key_points:
            st.markdown("### Summary")
            st.markdown(f'<div class="summary-box">{html.escape(st.session_state.document_summary).replace("\n", "<br>")}</div>', unsafe_allow_html=True)
            st.markdown("### Key Points")
            st.markdown(f'<div class="answer-box">{html.escape(st.session_state.document_key_points).replace("\n", "<br>")}</div>', unsafe_allow_html=True)
            down_l, down_r = st.columns(2)
            with down_l: st.download_button("Download Summary TXT", data=st.session_state.summary_report, file_name="summary_report.txt", mime="text/plain")
            with down_r: st.download_button("Download Summary DOCX", data=st.session_state.summary_docx, file_name="summary_report.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
elif st.session_state.active_view == "History":
    st.markdown('<div class="panel-header-title">Recent questions</div>', unsafe_allow_html=True)
    if st.session_state.query_history:
        for run_item in reversed(st.session_state.query_history):
            with st.expander(f"{run_item['time']} | {run_item['language']} | {run_item['question']}"):
                st.markdown("**Question**")
                st.write(run_item["question"])
                st.markdown("**Answer**")
                st.write(run_item["answer"])
    else:
        st.info("No questions yet.")
